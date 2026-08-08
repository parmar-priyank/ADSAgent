"""
routers/qc_checks.py — QC checklist execution and history routes.

Routes:
  GET  /qc-check
  POST /run-checklist
  GET  /checklist-result
  POST /checklist-analyse-file
  POST /checklist-save-edits
  GET  /checklist-download
  POST /checklist-confirm
  POST /checklist-save-draft
  GET  /user/history
  GET  /user/qc-version/{version_id}
  POST /user/qc-version/{version_id}/save
  POST /user/qc-version/{version_id}/add-email
  GET  /post-qc
  GET  /post-qc/customer/{quote_id}
"""
import asyncio
import base64
import io
import logging
import json
from concurrent.futures import ThreadPoolExecutor
import mimetypes
import os
import re
import zipfile as zipmod
from datetime import datetime

import pdfplumber

try:
    import fitz as _fitz   # PyMuPDF — optional; graceful fallback to text-only if missing
except ImportError:
    _fitz = None

try:
    import pillow_heif   # HEIC/HEIF support for Pillow — iPhone photos are often .heic
    pillow_heif.register_heif_opener()
except ImportError:
    pass

try:
    import docx as _docx   # python-docx — optional; .docx refs report a clear error if missing
except ImportError:
    _docx = None

_CLAUDE_IMAGE_MIMES = {"image/jpeg", "image/png", "image/gif", "image/webp"}

_PDF_DPI = 100        # 100 DPI is enough for Claude vision and cuts payload to 1/4
# Send at most this many pages per PDF to Claude. Was 2, which silently hid
# anything on page 3+ (e.g. a signature block on page 3 of a 4-page Signed
# Agreement) from every checklist item checking that PDF — no prompt wording
# could fix that, since the page was never sent. Post-QC documents (e.g.
# install/handover packs) run up to ~30 pages, so the cap covers that.
_PDF_MAX_PAGES = 30

import db.audit_repo as audit
import db.quote_repo as db
import db.checklist_repo as tdb
from db.checklist_repo import store_pending_result, fetch_pending_result, PENDING_TTL
from reports.xlsx_builder import build_xlsx
from db.connection import get_db

from fastapi import APIRouter, Depends, File, Form, HTTPException, Request, UploadFile
from fastapi.responses import HTMLResponse, RedirectResponse, Response
from itsdangerous import BadSignature

from config import (
    CLAUDE_MAX_CONCURRENCY,
    CLAUDE_MODEL,
    CLAUDE_PRICE_PER_M_INPUT,
    CLAUDE_PRICE_PER_M_OUTPUT,
    MAX_UPLOAD_BYTES,
    _index_context,
    _NO_CACHE,
    _get_claude,
    _resolve_theme,
    _signer,
    limiter,
    require_qc_access,
    templates,
)

router = APIRouter()
logger = logging.getLogger(__name__)

# Tracks which user IDs currently have a checklist run in progress.
# Prevents a user from submitting two overlapping jobs that would race
# on the same signed tokens. Each worker process has its own set, which
# is fine — jobs from the same user can land on different workers safely.
_active_jobs: set[int] = set()

# asyncio.to_thread always uses the loop's DEFAULT executor, which is capped
# at min(32, cpu_count + 4) — 6 threads on a 2-core droplet. That cap, not
# CPU or memory, is what made a 49-item checklist slow: the calls just queued.
# These helpers run the same work on a dedicated pool sized to
# CLAUDE_MAX_CONCURRENCY instead, with a semaphore so the ceiling is enforced
# rather than merely implied by the pool size.
_qc_executor = ThreadPoolExecutor(
    max_workers=CLAUDE_MAX_CONCURRENCY,
    thread_name_prefix="qc-claude",
)


def _bounded_gather(sem: asyncio.Semaphore, fn, *args):
    """Run a blocking fn on the QC pool, holding a concurrency slot."""
    async def _runner():
        async with sem:
            loop = asyncio.get_running_loop()
            return await loop.run_in_executor(_qc_executor, fn, *args)
    return _runner()


QC_SYSTEM = (
    "You are a QC document checker. "
    "Determine if the requirement is met based on the provided content. "
    'Reply with JSON only: {"status": "Yes" or "No", "remark": "one sentence explanation"}'
)


def _render_pdf(fname: str, fdata: bytes) -> tuple[str, list[bytes]]:
    """Return (extracted_text, [page_png_bytes]) for a PDF.

    Images are only rendered when text extraction yields nothing (scanned PDFs).
    Called once per unique PDF; the cache ensures each file is processed once.
    """
    text = ""
    try:
        with pdfplumber.open(io.BytesIO(fdata)) as pdf:
            pages_text = [p.extract_text() or "" for p in pdf.pages]
        text = "\n".join(pages_text)
    except Exception:
        pass

    # Always also render page images, not just when text extraction fails —
    # pdfplumber's extract_text() reads a PDF's content stream in whatever
    # order the text was drawn, which for table-heavy documents (e.g. an
    # insurance certificate with a label/value grid) can come out with every
    # label bunched up front and every value bunched after, unpaired and
    # useless despite being "non-empty" text. Sending the rendered page image
    # alongside that text lets Claude actually read the table visually
    # instead of only getting the scrambled text.
    pages: list[bytes] = []
    if _fitz is not None:
        try:
            doc = _fitz.open(stream=fdata, filetype="pdf")
            for page in doc[:_PDF_MAX_PAGES]:
                pix = page.get_pixmap(dpi=_PDF_DPI)
                pages.append(pix.tobytes("png"))
            doc.close()
        except Exception:
            pass

    return text, pages


_DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
_DOC_MIME  = "application/msword"   # legacy Word 97-2003; unreadable, but worth naming
_WORD_MIMES = (_DOCX_MIME, _DOC_MIME)


def _extract_docx_text(fname: str, fdata: bytes) -> str:
    """Return the readable text of a .docx (paragraphs + table cells).

    Claude's API takes images and PDFs, never .docx, so a Word file has to be
    reduced to text before it can be checked at all. Table cells are included
    because solar paperwork puts most of its real content (model numbers,
    serials, quantities) in tables, which paragraph text alone would miss.

    Note this is text only — images embedded in the Word file (signatures,
    photos, logos) are NOT visible to Claude this way, so an item asking
    "is it signed?" can't be answered from a .docx the way it can from a PDF
    or a photo. Raises on a corrupt/unreadable file so the caller can report
    a specific reason instead of silently checking an empty document.
    """
    if _docx is None:
        raise RuntimeError("python-docx is not installed on the server")
    document = _docx.Document(io.BytesIO(fdata))
    parts: list[str] = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _mime_of(fname: str, fbytes: bytes) -> str:
    mime, _ = mimetypes.guess_type(fname)
    if mime:
        return mime
    if fbytes[:4] == b"%PDF":
        return "application/pdf"
    if fbytes[:8] == b"\x89PNG\r\n\x1a\n":
        return "image/png"
    if fbytes[:3] == b"\xff\xd8\xff":
        return "image/jpeg"
    if fbytes[:6] in (b"GIF87a", b"GIF89a"):
        return "image/gif"
    if fbytes[:4] == b"RIFF" and fbytes[8:12] == b"WEBP":
        return "image/webp"
    if fbytes[:2] == b"BM":
        return "image/bmp"
    # A .docx is a ZIP containing word/document.xml — checked last so a real
    # ZIP archive isn't mistaken for a Word file, and only reached when the
    # filename gave us nothing (mimetypes already handles a normal .docx).
    if fbytes[:4] == b"PK\x03\x04":
        try:
            with zipmod.ZipFile(io.BytesIO(fbytes)) as _zf:
                if "word/document.xml" in _zf.namelist():
                    return _DOCX_MIME
        except Exception:
            pass
    return "application/octet-stream"


def _claude_safe_image(fdata: bytes, mime: str) -> tuple[bytes, str] | None:
    """
    Claude's vision API only accepts image/jpeg, image/png, image/gif,
    image/webp — phone photos are frequently .heic/.heif (Apple's default
    format), which mimetypes.guess_type reports correctly but Claude rejects
    outright with a 400. Convert those (and anything else unrecognized) to
    PNG via Pillow so every image ever reaches Claude in an accepted format.
    Returns (converted_bytes, "image/png") on success, or None if the image
    can't be decoded/converted at all (caller falls back to a clear error
    instead of sending bytes Claude will just reject again).
    """
    if mime in _CLAUDE_IMAGE_MIMES:
        return fdata, mime
    try:
        from PIL import Image
        img = Image.open(io.BytesIO(fdata))
        img = img.convert("RGB")
        out = io.BytesIO()
        img.save(out, format="PNG")
        return out.getvalue(), "image/png"
    except Exception:
        return None


def _parse_claude_json(raw: str) -> dict | None:
    raw = raw.strip()
    if raw.startswith("```"):
        raw = raw.split("```", 2)[1]
        if raw.startswith("json"):
            raw = raw[4:]
        raw = raw.rsplit("```", 1)[0].strip()
    try:
        return json.loads(raw)
    except Exception:
        return None


_QC_BATCH_SYSTEM = (
    "You are a QC document checker. You will be shown ONE reference document "
    "(as text and/or page images) followed by a numbered list of independent "
    "checklist items that all apply to that SAME document. Evaluate each item "
    "completely independently — the requirement, status, and remark for one "
    "item must not be influenced by any other item's requirement or answer. "
    'Reply with JSON only: a single array, one object per item in the same '
    'order given, each shaped exactly like '
    '{"item_index": <the item\'s given index number>, "status": "Yes" or "No" or "N/A", '
    '"remark": "one sentence explanation"}. Do not omit any item, do not add '
    "extra items, and do not wrap the array in any other object."
)


def _claude_check_batch(client, user_content, item_indices: list[int]) -> tuple[dict[int, dict], int, int]:
    """Same request/response contract as _claude_check, but for N checklist
    items that all share one reference file: one Claude call instead of N,
    since the expensive part of the request (the file's page images) would
    otherwise be re-uploaded once per item. user_content already contains the
    shared document content followed by each item's numbered requirement
    text (built by the caller); item_indices lists the item_index values
    used in that text, in order, so a response can be validated against
    exactly the set of items that were actually asked about.

    Returns ({item_index: {"status":..., "remark":...}}, input_tokens,
    output_tokens) — the two token counts are the TOTAL for this one shared
    call (summed across the retry too, if it fires), not split per item;
    the caller adds them once to its running total rather than distributing
    them across items and re-summing, which would be the same number with
    more code. On any failure (bad JSON twice, missing/extra indices,
    exception), every item_index maps to the same N/A fallback — callers
    must never assume a partial batch result, only all-or-nothing.
    """
    wanted = set(item_indices)
    messages = [{"role": "user", "content": user_content}]

    def _fallback(remark: str) -> dict[int, dict]:
        return {idx: {"status": "N/A", "remark": remark} for idx in item_indices}

    def _try_parse(raw: str):
        parsed = _parse_claude_json(raw)
        if not isinstance(parsed, list):
            return None
        by_idx = {}
        for entry in parsed:
            if not isinstance(entry, dict) or "item_index" not in entry:
                return None
            try:
                idx = int(entry["item_index"])
            except (TypeError, ValueError):
                return None
            by_idx[idx] = {"status": entry.get("status", "N/A"), "remark": entry.get("remark", "")}
        if set(by_idx.keys()) != wanted:
            # Claude dropped, duplicated, or invented an item_index — the
            # response can't be trusted to map back onto the right rows.
            return None
        return by_idx

    usage_in = usage_out = 0
    try:
        resp = client.messages.create(
            model=CLAUDE_MODEL, max_tokens=max(500, 150 * len(item_indices)),
            system=_QC_BATCH_SYSTEM, messages=messages,
        )
        usage_in  += getattr(resp.usage, "input_tokens", 0) or 0
        usage_out += getattr(resp.usage, "output_tokens", 0) or 0
        if usage_in == 0 and usage_out == 0:
            # See matching note in _claude_check — a real, successful Claude
            # response should never report 0/0 usage; log the raw usage
            # object so a recurrence is diagnosable instead of guessed at.
            logger.warning("Claude batch call returned real content but zero usage tokens; raw usage=%r", resp.usage)
        raw = resp.content[0].text if resp.content else ""
        result = _try_parse(raw)
        if result is None:
            # Same corrective-retry pattern as _claude_check, adapted for an
            # array reply — give the model one chance to fix its own format.
            messages.append({"role": "assistant", "content": raw})
            messages.append({"role": "user", "content":
                "That was not a valid response. Reply with ONLY a JSON array, one object "
                "per item, covering EXACTLY these item_index values and no others: "
                f"{sorted(wanted)}. Each object: "
                '{"item_index": <number>, "status": "Yes" or "No" or "N/A", "remark": "one sentence explanation"}'})
            resp2 = client.messages.create(
                model=CLAUDE_MODEL, max_tokens=max(500, 150 * len(item_indices)),
                system=_QC_BATCH_SYSTEM, messages=messages,
            )
            usage_in  += getattr(resp2.usage, "input_tokens", 0) or 0
            usage_out += getattr(resp2.usage, "output_tokens", 0) or 0
            raw2 = resp2.content[0].text if resp2.content else ""
            result = _try_parse(raw2)
            if result is None:
                logger.warning(
                    "Claude batch check returned unparseable/mismatched JSON twice "
                    "for item_indices=%s; raw replies: %r / %r", item_indices, raw, raw2,
                )
                return _fallback("AI returned an unreadable response for this batched check."), usage_in, usage_out
        return result, usage_in, usage_out
    except Exception as e:
        return _fallback(f"Error analysing batched file: {e}"), usage_in, usage_out


def _claude_check(client, user_content) -> dict:
    messages = [{"role": "user", "content": user_content}]
    resp = client.messages.create(
        model=CLAUDE_MODEL, max_tokens=500, system=QC_SYSTEM, messages=messages,
    )
    # Token accounting is additive to the return dict (private-prefixed keys),
    # so every existing caller that only reads "status"/"remark" is unaffected;
    # accumulated across both calls if the corrective retry below fires, since
    # both calls are real spend on this one checklist item.
    usage_in  = getattr(resp.usage, "input_tokens", 0) or 0
    usage_out = getattr(resp.usage, "output_tokens", 0) or 0
    if usage_in == 0 and usage_out == 0:
        # A real, successful Claude response should never report 0/0 usage —
        # log the raw usage object so a recurrence is diagnosable from the
        # server logs instead of guessed at (root cause not yet confirmed;
        # a fresh Pre-QC run on 2026-07-23 produced real Yes/No verdicts but
        # 0 total tokens with no exception raised anywhere in this chain).
        logger.warning("Claude call returned real content but zero usage tokens; raw usage=%r", resp.usage)
    raw = resp.content[0].text if resp.content else ""
    r = _parse_claude_json(raw)
    if r is None:
        # The model occasionally answers with prose instead of the requested
        # JSON (e.g. when a document's extracted text is jumbled/hard to
        # read) — one corrective retry recovers most of these; only give up
        # and log the raw reply if it fails twice, so a stuck item is
        # debuggable from the server logs instead of just "N/A" with no trace.
        messages.append({"role": "assistant", "content": raw})
        messages.append({"role": "user", "content":
            'That was not valid JSON. Reply with ONLY the JSON object, nothing else: '
            '{"status": "Yes", "No", or "N/A", "remark": "one sentence explanation"}'})
        resp2 = client.messages.create(
            model=CLAUDE_MODEL, max_tokens=500, system=QC_SYSTEM, messages=messages,
        )
        usage_in  += getattr(resp2.usage, "input_tokens", 0) or 0
        usage_out += getattr(resp2.usage, "output_tokens", 0) or 0
        raw2 = resp2.content[0].text if resp2.content else ""
        r = _parse_claude_json(raw2)
        if r is None:
            logger.warning("Claude returned unparseable JSON twice; raw replies: %r / %r", raw, raw2)
            return {"status": "N/A", "remark": "AI returned an unreadable response.",
                    "_input_tokens": usage_in, "_output_tokens": usage_out}
    return {"status": r.get("status", "N/A"), "remark": r.get("remark", ""),
            "_input_tokens": usage_in, "_output_tokens": usage_out}


# ---------------------------------------------------------------------------
# Email verify (Step 2) — upload .eml, extract attachments, match vs reference PDF
# ---------------------------------------------------------------------------

def _parse_eml(raw: bytes) -> list[dict]:
    """
    Parse a raw .eml file and return a list of attachment dicts:
      {name, mime, data, size_kb}
    Inline images and text parts are skipped — only named attachments are returned.
    Uses Python's stdlib email parser — no extra dependencies needed.
    """
    import email as _email
    import email.policy as _policy

    msg = _email.message_from_bytes(raw, policy=_policy.compat32)
    attachments = []
    for part in msg.walk():
        disposition = part.get_content_disposition() or ""
        fname = part.get_filename()
        if not fname:
            continue
        # Decode RFC-2047 encoded filenames (e.g. =?utf-8?q?...?=)
        from email.header import decode_header as _dh
        decoded_parts = _dh(fname)
        fname = "".join(
            (t.decode(enc or "utf-8") if isinstance(t, bytes) else t)
            for t, enc in decoded_parts
        )
        payload = part.get_payload(decode=True)
        if payload is None:
            continue
        attachments.append({
            "name": fname,
            "mime": part.get_content_type(),
            "data": payload,
            "size_kb": round(len(payload) / 1024, 1),
        })
    return attachments


def _match_attachment_with_claude(client, attachment: dict, reference_text: str) -> dict:
    """
    Send one email attachment to Claude with the reference PDF text and ask
    whether the attachment details match the signed agreement.
    Returns {name, size_kb, mime, status, remark}.
    """
    name = attachment["name"]
    data = attachment["data"]
    mime = attachment["mime"]
    size_kb = attachment["size_kb"]

    prompt = (
        f"You are verifying whether an email attachment sent to a customer "
        f"matches the details of their signed solar agreement.\n\n"
        f"Attachment filename: {name}\n\n"
        f"--- SIGNED AGREEMENT DETAILS ---\n{reference_text}\n\n"
        f"Based on the attachment content and the agreement details above, "
        f"answer: does this attachment appear to be correct and consistent "
        f"with the signed agreement? "
        f'Reply with JSON only: {{"status": "Yes" or "No" or "N/A", '
        f'"remark": "one concise sentence explaining your finding"}}'
    )

    user_content: list = [{"type": "text", "text": prompt}]

    # If it is a PDF, render pages and send as images
    if mime == "application/pdf" or name.lower().endswith(".pdf"):
        pages: list[bytes] = []
        if _fitz is not None:
            try:
                doc = _fitz.open(stream=data, filetype="pdf")
                for page in doc:
                    pix = page.get_pixmap(dpi=_PDF_DPI)
                    pages.append(pix.tobytes("png"))
                doc.close()
            except Exception:
                pass
        # Also extract text
        try:
            with pdfplumber.open(io.BytesIO(data)) as pdf:
                txt = "\n".join(p.extract_text() or "" for p in pdf.pages)
            if txt.strip():
                user_content.append({"type": "text", "text": f"Attachment text:\n{txt[:6000]}"})
        except Exception:
            pass
        for png in pages:
            b64 = base64.standard_b64encode(png).decode()
            user_content.append({"type": "image", "source": {"type": "base64", "media_type": "image/png", "data": b64}})
    elif mime.startswith("image/"):
        safe = _claude_safe_image(data, mime)
        if safe is None:
            return {"name": name, "size_kb": size_kb, "mime": mime,
                    "status": "N/A", "remark": "Could not read this image format — try re-saving as JPG or PNG."}
        safe_data, safe_mime = safe
        b64 = base64.standard_b64encode(safe_data).decode()
        user_content.append({"type": "image", "source": {"type": "base64", "media_type": safe_mime, "data": b64}})
    else:
        # Plain text / docx / xlsx — send raw text if decodable
        try:
            user_content.append({"type": "text", "text": f"Attachment content (raw text):\n{data.decode('utf-8', errors='replace')[:6000]}"})
        except Exception:
            return {"name": name, "size_kb": size_kb, "mime": mime,
                    "status": "N/A", "remark": "Unsupported attachment format — cannot read content."}

    try:
        resp = client.messages.create(
            model=CLAUDE_MODEL,
            max_tokens=300,
            messages=[{"role": "user", "content": user_content}],
        )
        usage_in  = getattr(resp.usage, "input_tokens", 0) or 0
        usage_out = getattr(resp.usage, "output_tokens", 0) or 0
        if usage_in == 0 and usage_out == 0:
            # See matching note in _claude_check — a real, successful Claude
            # response should never report 0/0 usage; log the raw usage
            # object so a recurrence is diagnosable instead of guessed at.
            logger.warning("Attachment-match Claude call returned real content but zero usage tokens; raw usage=%r", resp.usage)
        raw_resp = resp.content[0].text.strip()
        if raw_resp.startswith("```"):
            raw_resp = raw_resp.split("```", 2)[1]
            if raw_resp.startswith("json"):
                raw_resp = raw_resp[4:]
            raw_resp = raw_resp.rsplit("```", 1)[0].strip()
        r = json.loads(raw_resp)
        return {"name": name, "size_kb": size_kb, "mime": mime,
                "status": r.get("status", "N/A"), "remark": r.get("remark", ""),
                "_input_tokens": usage_in, "_output_tokens": usage_out}
    except Exception as e:
        return {"name": name, "size_kb": size_kb, "mime": mime,
                "status": "N/A", "remark": f"Error analysing attachment: {e}",
                "_input_tokens": 0, "_output_tokens": 0}


def _analyse_single_file(client, item_text: str, prompt_text: str, reference_pdf_text: str,
                         files: list[tuple[str, bytes]]) -> dict:
    """
    Analyse one or more uploaded files against a single checklist item's
    prompt — the per-row "Upload File" path in the result-page edit mode, so
    a user can fix one item's evidence without re-uploading the whole ZIP.
    Mirrors the relevant half of _analyse_item()'s body (run-checklist),
    minus the ZIP-lookup logic, since here the files are already in hand.

    Multiple non-.eml files are combined into ONE Claude call (same
    all-files-together approach _analyse_item already uses when an item's
    reference resolves to 2+ files), so e.g. a multi-page scan split across
    several PDFs/photos gets one combined verdict instead of one per file.
    .eml files are still handled through their own dedicated
    attachment-matching path, same as the existing single-file behavior.
    """
    ref_section = (
        f"\n\n--- MAIN REFERENCE PDF (Signed Agreement) ---\n{reference_pdf_text}"
        if reference_pdf_text else ""
    )
    context = f"Checklist item: {item_text}\nRequirement: {prompt_text}{ref_section}"

    try:
        eml_files    = [(fn, fd) for fn, fd in files if fn.lower().endswith(".eml")]
        other_files  = [(fn, fd) for fn, fd in files if not fn.lower().endswith(".eml")]

        if eml_files and other_files:
            return {"status": "N/A", "remark": "Please upload either .eml file(s) or document/image file(s), not a mix, for this item."}

        if eml_files:
            all_atts = []
            for fname, fdata in eml_files:
                atts = _parse_eml(fdata)
                if atts:
                    all_atts.extend(atts)
            if not all_atts:
                return {"status": "N/A", "remark": "No attachments found in the uploaded .eml file(s)."}
            att_results = [_match_attachment_with_claude(client, att, reference_pdf_text) for att in all_atts]
            no_hits  = [r for r in att_results if r.get("status") == "No"]
            yes_hits = [r for r in att_results if r.get("status") == "Yes"]
            if no_hits:
                names = ", ".join(r["name"] for r in no_hits)
                return {"status": "No", "remark": f"Mismatch in attachment(s): {names}. " + no_hits[0].get("remark", "")}
            if yes_hits:
                names = ", ".join(r["name"] for r in yes_hits)
                return {"status": "Yes", "remark": f"Attachment(s) match agreement: {names}."}
            return {"status": "N/A", "remark": att_results[0].get("remark", "Could not verify email attachments.")}

        user_content = [{"type": "text", "text": context}]
        unsupported: list[str] = []
        unreadable_images: list[str] = []

        for fname, fdata in other_files:
            mime = _mime_of(fname, fdata)
            if mime == "application/pdf":
                text, pages = _render_pdf(fname, fdata)
                if text.strip():
                    user_content.append({"type": "text", "text": f"Document text ({fname}):{text[:8000]}"})
                for page_png in pages[:_PDF_MAX_PAGES]:
                    b64 = base64.standard_b64encode(page_png).decode()
                    user_content.append({"type": "image", "source": {"type": "base64", "media_type": "image/png", "data": b64}})
            elif mime in _WORD_MIMES:
                # Text only — see _extract_docx_text; images inside the Word
                # file (signatures, photos) are not visible to Claude here.
                if fname.lower().endswith(".doc"):
                    return {"status": "N/A", "remark":
                            f"'{fname}' is a legacy Word 97-2003 file, which cannot be read. Re-save it as .docx or PDF."}
                if _docx is None:
                    return {"status": "N/A", "remark":
                            f"Could not read '{fname}'. Install python-docx on the server (pip install python-docx)."}
                try:
                    docx_text = _extract_docx_text(fname, fdata)
                except Exception as exc:
                    return {"status": "N/A", "remark": f"Could not read Word document '{fname}': {exc}"}
                if docx_text.strip():
                    user_content.append({"type": "text", "text": f"Document text ({fname}):{docx_text[:8000]}"})
            elif mime.startswith("image/"):
                safe = _claude_safe_image(fdata, mime)
                if safe is None:
                    unreadable_images.append(fname)
                    continue
                safe_data, safe_mime = safe
                b64 = base64.standard_b64encode(safe_data).decode()
                user_content.append({"type": "image", "source": {"type": "base64", "media_type": safe_mime, "data": b64}})
            else:
                unsupported.append(fname)

        if len(user_content) == 1:
            if unsupported:
                return {"status": "N/A", "remark": f"Unsupported file type: {', '.join(unsupported)}"}
            if unreadable_images:
                return {"status": "N/A", "remark": f"Could not read image format: {', '.join(unreadable_images)} — try re-saving as JPG or PNG."}
            return {"status": "N/A", "remark": "No readable content found in the uploaded file(s)."}

        if unsupported or unreadable_images:
            note = "; ".join(filter(None, [
                f"unsupported: {', '.join(unsupported)}" if unsupported else "",
                f"unreadable image(s): {', '.join(unreadable_images)}" if unreadable_images else "",
            ]))
            user_content.append({"type": "text", "text": f"(Note: some uploaded files could not be used — {note}.)"})

        return _claude_check(client, user_content)

    except Exception as e:
        return {"status": "N/A", "remark": f"Error analysing file: {e}"}


def _build_reference_text(quote_id) -> str:
    """Build the same 'Quote Number: ...\\nCustomer Name: ...' reference
    block used when matching an uploaded file against a quote's signed
    agreement — shared by the "add one more email" endpoints below."""
    if not quote_id:
        return ""
    try:
        ref_record = db.get_quote(int(quote_id))
    except (ValueError, TypeError):
        ref_record = db.find_by_quote_number(quote_id)
    if not ref_record:
        return ""
    fields = [
        ("Quote Number",    ref_record.get("quote_number", "")),
        ("Customer Name",   ref_record.get("customer_name", "")),
        ("Install Date",    ref_record.get("install_date", "")),
        ("System Price",    ref_record.get("system_price", "")),
        ("STC Incentive",   ref_record.get("stc_incentive", "")),
        ("VIC Rebate",      ref_record.get("vic_rebate", "")),
        ("Battery Rebate",  ref_record.get("battery_rebate", "")),
        ("Total Price",     ref_record.get("total_price", "")),
        ("Deposit",         ref_record.get("deposit", "")),
        ("Balance",         ref_record.get("balance", "")),
        ("Payment Terms",   ref_record.get("payment_terms", "")),
        ("Billing Address", ref_record.get("billing_address", "")),
        ("Delivery Address",ref_record.get("delivery_address", "")),
        ("Roof Type",       ref_record.get("roof_type", "")),
        ("Email",           ref_record.get("email", "")),
        ("Phone",           ref_record.get("phone", "")),
        ("Notes",           ref_record.get("notes", "")),
    ]
    reference_text = "\n".join(f"{k}: {v}" for k, v in fields if v)
    line_items = ref_record.get("line_items") or []
    if line_items:
        reference_text += "\n\nSystem Components:\n" + "\n".join(
            f"- {li.get('item','')}: {li.get('specification','')}"
            for li in line_items if li.get("item")
        )
    return reference_text


def _derive_email_summary(results: list) -> tuple:
    """From a list of analyzed attachment results, rebuild the "N emails
    uploaded" header panel's data: (emails_meta, email_from, email_subject).
    emails_meta is set (multi-email view) when 2+ distinct emails are
    represented; otherwise email_from/email_subject carry the single one
    (matches the original single-email upload response shape)."""
    seen, seen_keys = [], set()
    for r in results:
        key = (r.get("source_address") or "", r.get("source_subject") or "")
        if key in seen_keys:
            continue
        seen_keys.add(key)
        seen.append({"from": r.get("source_email", ""), "subject": r.get("source_subject", "")})
    if len(seen) > 1:
        return seen, None, None
    if seen:
        return None, seen[0]["from"], seen[0]["subject"]
    return None, None, None


@router.get("/email-verify", response_class=HTMLResponse)
def email_verify_page(request: Request, quote_id: str = "", user=Depends(require_qc_access)):
    if user.get("role") != "admin" and not user.get("can_pre_qc"):
        raise HTTPException(403, "You do not have Pre-QC access.")
    results = None
    emails_meta = None
    email_subject = None
    email_from = None

    draft_json = ""
    if quote_id:
        try:
            draft_json = db.get_draft_email_results(int(quote_id))
        except (ValueError, TypeError):
            draft_json = ""

    if draft_json.strip():
        try:
            parsed = json.loads(draft_json)
            if isinstance(parsed, list) and parsed:
                results = parsed
                emails_meta, email_from, email_subject = _derive_email_summary(results)
        except Exception:
            results = None

    resp = templates.TemplateResponse(request, "user_email_verify.html", {
        "current_user": user,
        "theme": _resolve_theme(user),
        "quote_id": quote_id,
        "results": results,
        "error": None,
        "email_subject": email_subject,
        "email_from": email_from,
        "emails": emails_meta,
        "from_draft": bool(results),
    })
    resp.headers.update(_NO_CACHE)
    return resp


def _decode_email_field(raw_val: str) -> str:
    """Decode RFC-2047 encoded header (e.g. =?Windows-1252?Q?...?=) to plain text."""
    from email.header import decode_header as _decode_header
    if not raw_val:
        return ""
    parts = _decode_header(raw_val)
    decoded = []
    for part, enc in parts:
        if isinstance(part, bytes):
            decoded.append(part.decode(enc or "utf-8", errors="replace"))
        else:
            decoded.append(part)
    return "".join(decoded).strip()


def _extract_email_address(from_header: str) -> str:
    """Pull just the bare address out of a From header for dedup matching."""
    import email.utils as _eutils
    _, addr = _eutils.parseaddr(from_header or "")
    return addr.strip().lower()


@router.post("/email-verify", response_class=HTMLResponse)
@limiter.limit("15/minute")
async def email_verify_post(
    request: Request,
    eml_files: list[UploadFile] = File(..., alias="eml_file"),
    quote_id: str = Form(""),
    user=Depends(require_qc_access),
):
    if user.get("role") != "admin" and not user.get("can_pre_qc"):
        raise HTTPException(403, "You do not have Pre-QC access.")
    def _err(msg: str):
        resp = templates.TemplateResponse(request, "user_email_verify.html", {
            "current_user": user,
            "theme": _resolve_theme(user),
            "quote_id": quote_id,
            "results": None,
            "error": msg,
            "email_subject": None,
            "email_from": None,
            "emails": None,
        })
        resp.headers.update(_NO_CACHE)
        return resp

    import email as _email_mod
    import email.policy as _policy

    # Build reference text from the linked quote (shared across all emails)
    reference_text = ""
    if quote_id:
        try:
            ref_record = db.get_quote(int(quote_id))
        except (ValueError, TypeError):
            ref_record = db.find_by_quote_number(quote_id)
        if ref_record:
            fields = [
                ("Quote Number",    ref_record.get("quote_number", "")),
                ("Customer Name",   ref_record.get("customer_name", "")),
                ("Install Date",    ref_record.get("install_date", "")),
                ("System Price",    ref_record.get("system_price", "")),
                ("Total Price",     ref_record.get("total_price", "")),
                ("Deposit",         ref_record.get("deposit", "")),
                ("Balance",         ref_record.get("balance", "")),
                ("Payment Terms",   ref_record.get("payment_terms", "")),
                ("Billing Address", ref_record.get("billing_address", "")),
                ("Email",           ref_record.get("email", "")),
                ("Phone",           ref_record.get("phone", "")),
            ]
            reference_text = "\n".join(f"{k}: {v}" for k, v in fields if v)

    client = _get_claude()

    emails_meta = []   # [{subject, from, address}] — one per uploaded .eml
    all_attachments = []  # flattened, each tagged with its source email index
    for idx, eml_file in enumerate(eml_files):
        if not eml_file.filename.lower().endswith(".eml"):
            return _err(f"'{eml_file.filename}' is not a valid .eml file (Outlook email format).")

        raw = await eml_file.read()
        if len(raw) > 50 * 1024 * 1024:  # 50 MB cap per email file
            return _err(f"'{eml_file.filename}' is too large. Maximum allowed size is 50 MB per email.")

        msg = _email_mod.message_from_bytes(raw, policy=_policy.compat32)
        subject = _decode_email_field(msg.get("Subject", "")) or "(no subject)"
        from_hdr = _decode_email_field(msg.get("From", "")) or "(unknown sender)"
        emails_meta.append({
            "subject": subject,
            "from": from_hdr,
            "address": _extract_email_address(from_hdr),
        })

        attachments = _parse_eml(raw)
        for att in attachments:
            all_attachments.append((idx, att))

    if not all_attachments:
        noun = "email" if len(eml_files) == 1 else "emails"
        return _err(f"No attachments found in the uploaded {noun}. Please check the .eml file(s).")

    # Analyse all attachments (across all emails) concurrently, capped at
    # CLAUDE_MAX_CONCURRENCY — an email with many attachments would otherwise
    # queue behind the 6-thread default executor exactly like a checklist run.
    _eml_sem = asyncio.Semaphore(CLAUDE_MAX_CONCURRENCY)
    analysed = await asyncio.gather(
        *[_bounded_gather(_eml_sem, _match_attachment_with_claude, client, att, reference_text)
          for _, att in all_attachments]
    )
    # Tokens actually spent by THIS upload round, added to the quote's
    # running (not-yet-consumed) Verify Email total — a user may verify
    # emails across several separate /email-verify submissions before ever
    # clicking "Continue to Upload ZIP", and each submission re-renders this
    # page fresh, so accumulating in the DB (rather than only in this
    # request's hidden field) is what keeps earlier rounds from being lost.
    # Travels forward via a hidden form field (same mechanism as
    # email_results_json); reset to 0 once a checklist run consumes it.
    email_verify_input_tokens = email_verify_output_tokens = 0
    if quote_id:
        try:
            qid_int = int(quote_id)
        except (ValueError, TypeError):
            qid_int = None
        if qid_int is not None:
            round_in  = sum(r.get("_input_tokens", 0) or 0 for r in analysed)
            round_out = sum(r.get("_output_tokens", 0) or 0 for r in analysed)
            email_verify_input_tokens, email_verify_output_tokens = \
                db.add_draft_email_verify_tokens(qid_int, round_in, round_out)

    results = []
    for (idx, _att), r in zip(all_attachments, analysed):
        results.append({
            **r,
            "ai_status": r.get("status"),
            "source_email": emails_meta[idx]["from"],
            "source_subject": emails_meta[idx]["subject"],
            "source_address": emails_meta[idx]["address"],
        })

    # Dedup exact duplicates: same sender address AND same attachment name.
    # Different attachments from the same sender (e.g. contract vs warranty
    # doc in separate emails) are kept — only true re-uploads collapse into
    # a single unified answer.
    seen_keys: set = set()
    deduped = []
    for r in results:
        key = ((r.get("source_address") or ""), (r.get("name") or "").strip().lower())
        if key[0] and key[1] and key in seen_keys:
            continue
        if key[0] and key[1]:
            seen_keys.add(key)
        deduped.append(r)
    results = deduped

    # Merge with any previously-saved draft for this quote, so uploading one
    # more .eml doesn't discard emails already analysed in an earlier visit —
    # new results replace old ones with the same (sender, attachment name),
    # everything else from the saved draft is kept alongside the new results.
    if quote_id:
        try:
            draft_json = db.get_draft_email_results(int(quote_id))
        except (ValueError, TypeError):
            draft_json = ""
        if draft_json.strip():
            try:
                saved = json.loads(draft_json)
                if isinstance(saved, list):
                    new_keys = {
                        ((r.get("source_address") or ""), (r.get("name") or "").strip().lower())
                        for r in results
                    }
                    kept_old = [
                        r for r in saved
                        if ((r.get("source_address") or ""), (r.get("name") or "").strip().lower())
                        not in new_keys
                    ]
                    results = kept_old + results
            except Exception:
                pass

    emails_meta, email_from, email_subject = _derive_email_summary(results)

    resp = templates.TemplateResponse(request, "user_email_verify.html", {
        "current_user": user,
        "theme": _resolve_theme(user),
        "quote_id": quote_id,
        "results": list(results),
        "error": None,
        "email_subject": email_subject,
        "email_from": email_from,
        "emails": emails_meta,
        "email_verify_input_tokens": email_verify_input_tokens,
        "email_verify_output_tokens": email_verify_output_tokens,
    })
    resp.headers.update(_NO_CACHE)
    return resp


@router.post("/email-verify/save-draft")
def email_verify_save_draft(
    quote_id: str = Form(""),
    email_results_json: str = Form(""),
    user=Depends(require_qc_access),
):
    """Persist Step 2's analyzed attachment results against the quote, so the
    user doesn't have to re-upload the same .eml file(s) if they come back
    to this quote later — independent of Continue/proceeding to Step 3."""
    if not quote_id:
        raise HTTPException(400, "Missing quote_id.")
    try:
        qid = int(quote_id)
    except (ValueError, TypeError):
        raise HTTPException(400, "Invalid quote_id.")

    try:
        parsed = json.loads(email_results_json) if email_results_json.strip() else []
        if not isinstance(parsed, list):
            parsed = []
    except Exception:
        parsed = []

    db.save_draft_email_results(qid, json.dumps(parsed))
    return {"ok": True}


# ---------------------------------------------------------------------------
# QC check page
# ---------------------------------------------------------------------------

@router.get("/qc-check", response_class=HTMLResponse)
def qc_check_page(request: Request, quote_id: str = "", kind: str = "pre",
                   version_id: str = "", user=Depends(require_qc_access)):
    kind = kind if kind in ("pre", "post") else "pre"
    if user.get("role") != "admin":
        if kind == "post" and not user.get("can_post_qc"):
            raise HTTPException(403, "You do not have Post-QC access.")
        if kind == "pre" and not user.get("can_pre_qc"):
            raise HTTPException(403, "You do not have Pre-QC access.")
    saved = tdb.list_templates(kind=kind)
    # version_id only ever comes from the admin "Upload ZIP" re-run action on
    # an existing Post-QC version — re-validated the same way upload_zip
    # does before being trusted, so a forged query string can't be used to
    # overwrite an unrelated version via this pre-fill. Template isn't
    # pre-selected: qc_versions only stores the template's NAME, not a
    # stable id, so admin just picks it again on this screen (same template
    # they're already replacing results for).
    prefill_version_id = ""
    if version_id.strip() and kind == "post" and user.get("role") == "admin":
        try:
            vid_int = int(version_id)
        except (ValueError, TypeError):
            vid_int = None
        if vid_int is not None:
            with get_db() as conn:
                row = conn.execute(
                    "SELECT id, quote_id, kind FROM qc_versions WHERE id = ?", (vid_int,)
                ).fetchone()
            existing = dict(row) if row else None
            if (existing and str(existing.get("quote_id")) == str(quote_id)
                    and (existing.get("kind") or "pre") == "post"):
                prefill_version_id = str(vid_int)
    resp = templates.TemplateResponse(request, "user_qc.html", {
        "current_user": user,
        "theme": _resolve_theme(user),
        "quote_id": quote_id,
        "kind": kind,
        "saved_templates": saved,
        "email_results_json": "",
        "error": None,
        "prefill_version_id": prefill_version_id,
    })
    resp.headers.update(_NO_CACHE)
    return resp


@router.post("/qc-check-with-email", response_class=HTMLResponse)
def qc_check_page_with_email(
    request: Request,
    quote_id: str = Form(""),
    email_results_json: str = Form(""),
    email_verify_input_tokens: int = Form(0),
    email_verify_output_tokens: int = Form(0),
    user=Depends(require_qc_access),
):
    """Arrive from Step 2 (email verify) carrying the edited email results.
    Always Pre-QC — Post-QC's flow never goes through email verify (see
    /post-qc/customer/{quote_id}, which links straight to /qc-check?kind=post)."""
    if user.get("role") != "admin" and not user.get("can_pre_qc"):
        raise HTTPException(403, "You do not have Pre-QC access.")
    saved = tdb.list_templates(kind="pre")
    resp = templates.TemplateResponse(request, "user_qc.html", {
        "current_user": user,
        "theme": _resolve_theme(user),
        "quote_id": quote_id,
        "kind": "pre",
        "saved_templates": saved,
        "email_results_json": email_results_json,
        "email_verify_input_tokens": email_verify_input_tokens,
        "email_verify_output_tokens": email_verify_output_tokens,
        "error": None,
    })
    resp.headers.update(_NO_CACHE)
    return resp


# ---------------------------------------------------------------------------
# Run checklist
# ---------------------------------------------------------------------------

@router.post("/run-checklist", response_class=HTMLResponse)
@limiter.limit("10/minute")
async def upload_zip(
    request: Request,
    zip_file: UploadFile = File(...),
    quote_id: str = Form(""),
    template_id: int = Form(...),
    kind: str = Form("pre"),
    email_results_json: str = Form(""),
    email_verify_input_tokens: int = Form(0),
    email_verify_output_tokens: int = Form(0),
    version_id: str = Form(""),
    user=Depends(require_qc_access),
):
    kind = kind if kind in ("pre", "post") else "pre"
    # Admin-only "Upload ZIP" re-run from an existing Post-QC version's
    # checklist modal — re-analyzes everything against a fresh ZIP but
    # OVERWRITES that same version in place instead of creating a new one.
    # Only ever valid for kind == "post": re-running Pre-QC this way isn't
    # offered anywhere in the UI, and validating here (not trusting the
    # form value alone) stops a stray/forged version_id from letting this
    # silently overwrite an unrelated customer's QC version.
    target_version_id = None
    if version_id.strip() and kind == "post" and user.get("role") == "admin":
        try:
            target_version_id = int(version_id)
        except (ValueError, TypeError):
            target_version_id = None
        if target_version_id is not None:
            existing = None
            with get_db() as conn:
                row = conn.execute(
                    "SELECT id, quote_id, kind FROM qc_versions WHERE id = ?",
                    (target_version_id,),
                ).fetchone()
                existing = dict(row) if row else None
            same_quote = existing and str(existing.get("quote_id")) == str(quote_id)
            is_post_version = existing and (existing.get("kind") or "pre") == "post"
            if not existing or not same_quote or not is_post_version:
                target_version_id = None

    if user.get("role") != "admin":
        if kind == "post" and not user.get("can_post_qc"):
            raise HTTPException(403, "You do not have Post-QC access.")
        if kind == "pre" and not user.get("can_pre_qc"):
            raise HTTPException(403, "You do not have Pre-QC access.")
    uid = user["id"]
    if uid in _active_jobs:
        saved = tdb.list_templates(kind=kind)
        resp = templates.TemplateResponse(request, "user_qc.html", {
            "current_user": user,
            "theme": _resolve_theme(user),
            "quote_id": quote_id,
            "kind": kind,
            "saved_templates": saved,
            "error": "A checklist run is already in progress for your account. Please wait for it to finish.",
        })
        resp.headers.update(_NO_CACHE)
        return resp

    tpl = tdb.get_template(template_id)
    if not tpl:
        raise HTTPException(404, "Template not found.")
    items = tdb.get_items(template_id)

    data = await zip_file.read()

    def _qc_error(error: str):
        saved = tdb.list_templates(kind=kind)
        resp = templates.TemplateResponse(request, "user_qc.html", {
            "current_user": user,
            "theme": _resolve_theme(user),
            "quote_id": quote_id,
            "kind": kind,
            "saved_templates": saved,
            "email_results_json": email_results_json,
            "error": error,
        })
        resp.headers.update(_NO_CACHE)
        return resp

    if len(data) > MAX_UPLOAD_BYTES:
        return _qc_error("ZIP file is too large. Maximum allowed size is 200 MB.")

    _ZIP_ENTRY_MAX = 500 * 1024 * 1024
    zip_files: dict[str, bytes] = {}
    duplicate_basenames: dict[str, list[str]] = {}
    try:
        with zipmod.ZipFile(io.BytesIO(data)) as zf:
            total_uncompressed = sum(e.file_size for e in zf.infolist())
            if total_uncompressed > _ZIP_ENTRY_MAX:
                return _qc_error("ZIP contents are too large. Total uncompressed size must not exceed 500 MB.")
            for name in zf.namelist():
                basename = os.path.basename(name).strip()
                if not basename:
                    continue
                key = basename.lower()
                if key in zip_files:
                    # Same filename appears in more than one folder inside the
                    # ZIP — matching only ever looks at basenames, so which one
                    # "wins" would be arbitrary (zip entry order, not anything
                    # meaningful). Rather than silently picking one and having
                    # the wrong file get analysed, surface this and let the
                    # user fix the ZIP instead of guessing.
                    duplicate_basenames.setdefault(key, [name]).append(name)
                    continue
                zip_files[key] = zf.read(name)
    except zipmod.BadZipFile:
        return _qc_error("Uploaded file is not a valid ZIP archive.")

    if duplicate_basenames:
        lines = "; ".join(
            f'"{os.path.basename(paths[0])}" appears at: {", ".join(paths)}'
            for paths in duplicate_basenames.values()
        )
        return _qc_error(
            "This ZIP has the same filename in more than one folder, so it's "
            "ambiguous which one to check: " + lines +
            ". Please rename or remove the duplicate and re-upload."
        )

    # Build reference context from the linked quote record
    reference_pdf_text = ""
    job_line_items: list = []
    if quote_id:
        try:
            ref_record = db.get_quote(int(quote_id))
        except (ValueError, TypeError):
            ref_record = db.find_by_quote_number(quote_id)
        if ref_record:
            fields = [
                ("Quote Number",    ref_record.get("quote_number", "")),
                ("Customer Name",   ref_record.get("customer_name", "")),
                ("Install Date",    ref_record.get("install_date", "")),
                ("System Price",    ref_record.get("system_price", "")),
                ("STC Incentive",   ref_record.get("stc_incentive", "")),
                ("VIC Rebate",      ref_record.get("vic_rebate", "")),
                ("Battery Rebate",  ref_record.get("battery_rebate", "")),
                ("Total Price",     ref_record.get("total_price", "")),
                ("Deposit",         ref_record.get("deposit", "")),
                ("Balance",         ref_record.get("balance", "")),
                ("Payment Terms",   ref_record.get("payment_terms", "")),
                ("Billing Address", ref_record.get("billing_address", "")),
                ("Delivery Address",ref_record.get("delivery_address", "")),
                ("Roof Type",       ref_record.get("roof_type", "")),
                ("Email",           ref_record.get("email", "")),
                ("Phone",           ref_record.get("phone", "")),
                ("Notes",           ref_record.get("notes", "")),
            ]
            reference_pdf_text = "\n".join(f"{k}: {v}" for k, v in fields if v)
            job_line_items = ref_record.get("line_items") or []
            if job_line_items:
                reference_pdf_text += "\n\nSystem Components:\n" + "\n".join(
                    f"- {li.get('item','')}: {li.get('specification','')}"
                    for li in job_line_items if li.get("item")
                )

    # Battery-only job detection — free, deterministic, no extra Claude call.
    # The original PDF extraction already captures every System/pricing table
    # row into line_items (see EXTRACTION_SCHEMA in services/ai_service.py),
    # including a "Panel"/"Panels" row whenever the job has solar panels.
    # Real production data confirms a battery-only job's line_items simply
    # has no Panel/Panels row at all (never an empty-quantity row), so its
    # absence is a reliable signal here — nothing inferred from wording/AI.
    # No line_items at all (e.g. quote_id missing) means "can't tell" —
    # treated as a panel job (run everything, as today).
    #
    # Pre-QC only: this gating exists to save Claude cost on Pre-QC runs,
    # which check the signed agreement itself (where "no Panel line item"
    # reliably means a battery-only job). Post-QC checks physical installer
    # evidence (photos/paperwork of the completed install) — an unrelated
    # document set where this signal doesn't apply — so it must always run
    # every item regardless of what the Pre-QC line_items looked like.
    is_battery_only_job = kind == "pre" and bool(job_line_items) and not any(
        (li.get("item") or "").strip().lower() in ("panel", "panels")
        for li in job_line_items
    )

    _claude_client = _get_claude()

    def _resolve_file(name: str):
        key = name.lower()
        if (fdata := zip_files.get(key)) is not None:
            return key, fdata

        key_nospace  = key.replace(" ", "")
        stem         = os.path.splitext(key)[0]
        stem_nospace = stem.replace(" ", "")

        best_name, best_data = None, None
        for zname, zbytes in zip_files.items():
            if zname.replace(" ", "") == key_nospace:
                return zname, zbytes
            zstem = os.path.splitext(zname)[0]
            if zstem == stem:
                best_name, best_data = zname, zbytes
            elif zstem.replace(" ", "") == stem_nospace and best_data is None:
                best_name, best_data = zname, zbytes

        return (best_name, best_data) if best_data is not None else (name, None)

    def _expected_kind(ref_name: str) -> str:
        """Guess what kind of file a reference name implies, from its extension."""
        ext = os.path.splitext(ref_name.lower())[1]
        if ext == ".eml":
            return "eml"
        if ext == ".pdf":
            return "pdf"
        if ext in (".docx", ".doc"):
            return "docx"
        if ext in (".jpg", ".jpeg", ".png", ".gif", ".webp", ".bmp"):
            return "image"
        return "other"

    def _actual_kind(fname: str, fdata: bytes) -> str:
        if fname.lower().endswith(".eml"):
            return "eml"
        mime = _mime_of(fname, fdata)
        if mime == "application/pdf":
            return "pdf"
        if mime in _WORD_MIMES:
            return "docx"
        if mime.startswith("image/"):
            return "image"
        return "other"

    # ── Resolve every checklist item's reference file(s) up front, sequentially ──
    # (must happen before the concurrent AI analysis below, so two items can't
    # both claim the same fallback-matched file in a race).
    #
    # Pass 1: exact-name matching (same as before) — claims files as it goes so
    # pass 2 knows what's already spoken for.
    claimed: set[str] = set()
    item_resolved: dict[int, list[tuple]] = {}

    for i, item in enumerate(items):
        if item.get("is_section"):
            continue
        ref = (item.get("reference") or "").strip()
        # "+" and "/" are both accepted as "all of these files together" —
        # some templates (typed directly, not via the Reference modal, which
        # always writes "+") use "/" instead, e.g. "job pack/installer
        # presence". Without this, the whole string was looked up as one
        # literal filename, which never matches anything in the ZIP and
        # silently reports "File not found in ZIP: job pack/installer
        # presence" even when both files are genuinely present.
        ref_names = [r.strip() for r in re.split(r"[+/]", ref) if r.strip()]
        resolved = []
        for rname in ref_names:
            zname, zdata = _resolve_file(rname)
            if zdata is not None:
                claimed.add(zname)
            resolved.append([rname, zname, zdata])
        item_resolved[i] = resolved

    # Pass 2: for references still unmatched, fall back to the single unclaimed
    # file left in the ZIP of the same kind (image/pdf/eml) the reference name
    # implies — only when that leaves exactly one candidate, so it's unambiguous.
    for i, resolved in item_resolved.items():
        for entry in resolved:
            rname, zname, zdata = entry
            if zdata is not None:
                continue
            kind = _expected_kind(rname)
            if kind == "other":
                continue
            candidates = [
                (fn, fd) for fn, fd in zip_files.items()
                if fn not in claimed and _actual_kind(fn, fd) == kind
            ]
            if len(candidates) == 1:
                fn, fd = candidates[0]
                claimed.add(fn)
                entry[1], entry[2] = fn, fd

    # Pre-render every REFERENCED PDF in the ZIP exactly once (cache avoids
    # re-rendering per item), concurrently across threads — same pattern the
    # AI analysis step below already uses. This used to loop over every file
    # in the ZIP sequentially and render all of them regardless of whether
    # any checklist item actually referenced them; rendering (pdfplumber text
    # extraction + full-page rasterization) is real, blocking CPU work per
    # PDF, so a ZIP with several unrelated/unreferenced PDFs (or just a few
    # multi-page ones) made every upload slow before any AI call even
    # started. Now only the files item_resolved actually points at get
    # rendered, and all of them render in parallel instead of one at a time.
    referenced_pdf_names = sorted({
        zname
        for resolved in item_resolved.values()
        for _, zname, zdata in resolved
        if zname is not None and zdata is not None and _mime_of(zname, zdata) == "application/pdf"
    })
    # Deliberately still on the DEFAULT executor, not the larger QC pool the
    # Claude calls use. Rendering is CPU-bound, so on a 2-core box extra
    # threads win nothing here and would only compete with the API calls for
    # the GIL. The Claude calls are the opposite — blocked on the network —
    # which is why only those were given a bigger pool.
    pdf_render_results = await asyncio.gather(
        *[asyncio.to_thread(_render_pdf, fname, zip_files[fname]) for fname in referenced_pdf_names]
    )
    pdf_cache: dict[str, tuple[str, list[bytes]]] = dict(zip(referenced_pdf_names, pdf_render_results))

    def _analyse_eml_item(eml_parts: list, context: str) -> dict:
        """
        Roll up a checklist item whose reference resolves to one or more .eml
        files inside the ZIP into a single Yes/No/N/A verdict, by extracting
        each email's attachments and matching them against the reference
        text — same logic the standalone "Verify Email" step (Step 2) uses.
        """
        all_atts = []
        for fname, fdata in eml_parts:
            atts = _parse_eml(fdata)
            if not atts:
                continue
            all_atts.extend(atts)

        if not all_atts:
            return {"status": "N/A", "remark": "No attachments found in the referenced .eml file(s)."}

        att_results = [
            _match_attachment_with_claude(_claude_client, att, reference_pdf_text)
            for att in all_atts
        ]
        # Every attachment match is its own Claude call — sum all of them so
        # the checklist item's own token count reflects the full cost of
        # verifying every attachment, not just whichever one wins the verdict.
        eml_in  = sum(r.get("_input_tokens", 0) or 0 for r in att_results)
        eml_out = sum(r.get("_output_tokens", 0) or 0 for r in att_results)

        no_hits  = [r for r in att_results if r.get("status") == "No"]
        yes_hits = [r for r in att_results if r.get("status") == "Yes"]
        if no_hits:
            names = ", ".join(r["name"] for r in no_hits)
            return {"status": "No", "remark": f"Mismatch in attachment(s): {names}. " + no_hits[0].get("remark", ""),
                    "_input_tokens": eml_in, "_output_tokens": eml_out}
        if yes_hits:
            names = ", ".join(r["name"] for r in yes_hits)
            return {"status": "Yes", "remark": f"Attachment(s) match agreement: {names}.",
                    "_input_tokens": eml_in, "_output_tokens": eml_out}
        return {"status": "N/A", "remark": att_results[0].get("remark", "Could not verify email attachments."),
                "_input_tokens": eml_in, "_output_tokens": eml_out}

    def _analyse_item(i: int, item: dict) -> dict:
        ref         = (item.get("reference") or "").strip()
        prompt_text = (item.get("prompt") or "").strip()

        if is_battery_only_job and not item.get("battery_only"):
            return {"status": "N/A", "remark": "Not applicable — battery-only job (no solar panels detected)."}
        if not ref:
            return {"status": "N/A", "remark": "No reference file specified."}
        if not prompt_text:
            return {"status": "N/A", "remark": "No prompt defined for this item."}

        # resolved entries are [ref_name, matched_zip_name_or_None, data_or_None],
        # already computed up front in the sequential resolve-and-claim pass above
        # (both the exact-name match and the same-file-type fallback).
        resolved_entries = item_resolved.get(i, [])
        missing = [rname for rname, zname, zdata in resolved_entries if zdata is None]
        found   = [(zname, zdata) for rname, zname, zdata in resolved_entries if zdata is not None]
        if not found:
            return {"status": "N/A", "remark": f"File not found in ZIP: {', '.join(missing)}"}
        resolved = found
        missing_note = (
            f" (Note: {', '.join(missing)} not found in ZIP, working with available files only.)"
            if missing else ""
        )

        ref_section = (
            f"\n\n--- MAIN REFERENCE PDF (Signed Agreement) ---\n{reference_pdf_text}"
            if reference_pdf_text else ""
        )
        context = f"Checklist item: {item['text']}\nRequirement: {prompt_text}{ref_section}{missing_note}"

        try:
            eml_parts, pdf_parts, image_parts, docx_parts, unsupported = [], [], [], [], []
            for fname, fdata in resolved:
                if fname.lower().endswith(".eml"):
                    eml_parts.append((fname, fdata))
                    continue
                mime = _mime_of(fname, fdata)
                if mime == "application/pdf":
                    pdf_parts.append((fname, fdata))
                elif mime in _WORD_MIMES:
                    docx_parts.append((fname, fdata))
                elif mime.startswith("image/"):
                    image_parts.append((fname, fdata, mime))
                else:
                    unsupported.append(fname)

            if eml_parts:
                return _analyse_eml_item(eml_parts, context)

            if unsupported:
                return {"status": "N/A", "remark": f"Unsupported file type: {', '.join(unsupported)}"}

            combined_pdf_text  = ""
            pdf_page_images: list[bytes] = []
            for fname, _ in pdf_parts:
                cached_text, cached_pages = pdf_cache.get(fname, ("", []))
                if cached_text.strip():
                    combined_pdf_text += f"\n\n--- {fname} ---\n{cached_text}"
                pdf_page_images.extend(cached_pages[:_PDF_MAX_PAGES])
                if not cached_pages and not cached_text.strip():
                    if _fitz is None:
                        return {"status": "N/A", "remark": f"Could not render '{fname}' as image. Install PyMuPDF (pip install pymupdf)."}

            # Word documents contribute text only (see _extract_docx_text) —
            # folded into the same block the PDF text uses, so an item whose
            # reference is a .docx takes the existing text-only Claude path.
            for fname, fdata in docx_parts:
                if fname.lower().endswith(".doc"):
                    return {"status": "N/A", "remark":
                            f"'{fname}' is a legacy Word 97-2003 file, which cannot be read. Re-save it as .docx or PDF."}
                if _docx is None:
                    return {"status": "N/A", "remark":
                            f"Could not read '{fname}'. Install python-docx on the server (pip install python-docx)."}
                try:
                    docx_text = _extract_docx_text(fname, fdata)
                except Exception as exc:
                    return {"status": "N/A", "remark": f"Could not read Word document '{fname}': {exc}"}
                if docx_text.strip():
                    combined_pdf_text += f"\n\n--- {fname} ---\n{docx_text}"

            user_content = [{"type": "text", "text": context}]
            if combined_pdf_text.strip():
                user_content.append({"type": "text", "text": f"Document text (for reference):{combined_pdf_text[:8000]}"})
            for page_png in pdf_page_images:
                b64 = base64.standard_b64encode(page_png).decode()
                user_content.append({"type": "image", "source": {"type": "base64", "media_type": "image/png", "data": b64}})
            unreadable_images = []
            for fname, fdata, mime in image_parts:
                safe = _claude_safe_image(fdata, mime)
                if safe is None:
                    unreadable_images.append(fname)
                    continue
                safe_data, safe_mime = safe
                b64 = base64.standard_b64encode(safe_data).decode()
                user_content.append({"type": "image", "source": {"type": "base64", "media_type": safe_mime, "data": b64}})

            if len(user_content) == 1:
                if unreadable_images:
                    return {"status": "N/A", "remark": f"Could not read image format: {', '.join(unreadable_images)} — try re-saving as JPG or PNG."}
                if combined_pdf_text.strip():
                    return _claude_check(_claude_client, f"{context}\n\nDocument content:{combined_pdf_text[:8000]}")
                return {"status": "N/A", "remark": "No readable content found in reference file(s)."}

            return _claude_check(_claude_client, user_content)

        except Exception as e:
            return {"status": "N/A", "remark": f"Error analysing file: {e}"}

    non_section = [(i, item) for i, item in enumerate(items) if not item.get("is_section")]

    # ── Batch items that share one reference file into a single Claude call,
    # instead of one call per item — applies to both Pre-QC and Post-QC.
    #
    # Eligible for batching: item_resolved[i] has exactly one entry, that
    # entry actually resolved to a file (zdata is not None), and that file's
    # kind is pdf/image (never .eml — email items keep their own dedicated
    # _analyse_eml_item path) AND the SAME file is the sole resolved
    # reference for 2+ items. A file referenced by only one item gets no
    # benefit from batching, so it stays on the plain per-item path too.
    batch_groups: dict[str, list[int]] = {}
    for i, item in non_section:
        if is_battery_only_job and not item.get("battery_only"):
            # Same gate as _analyse_item's early return — keep this item off
            # the paid batch path entirely on a battery-only job, since its
            # result would just be discarded in favor of a fixed N/A anyway.
            continue
        resolved_entries = item_resolved.get(i, [])
        if len(resolved_entries) != 1:
            continue
        rname, zname, zdata = resolved_entries[0]
        if zdata is None or zname is None:
            continue
        if zname.lower().endswith(".eml"):
            continue
        if _actual_kind(zname, zdata) not in ("pdf", "image"):
            continue
        if not (item.get("prompt") or "").strip():
            # No prompt defined — _analyse_item already returns a fixed
            # N/A for this without calling Claude at all; keep that exact
            # zero-cost behavior instead of pulling it into a paid batch call.
            continue
        batch_groups.setdefault(zname, []).append(i)
    batched_item_indices = {
        i for idxs in batch_groups.values() if len(idxs) >= 2 for i in idxs
    }

    def _build_batch_user_content(zname: str, idxs: list[int]) -> list:
        """Same content shape _analyse_item builds for a single item (text
        context + document text + page images), but with one shared
        document section followed by every batched item's own numbered
        requirement — so the model still answers each item from its own
        exact prompt text, just without re-uploading the file per item."""
        fdata = zip_files[zname]
        mime = _actual_kind(zname, fdata)

        items_block_lines = []
        for i in idxs:
            item = items[i]
            items_block_lines.append(
                f"item_index {i}: Checklist item: {item['text']}\nRequirement: {(item.get('prompt') or '').strip()}"
            )
        items_block = "\n\n".join(items_block_lines)

        ref_section = (
            f"\n\n--- MAIN REFERENCE PDF (Signed Agreement) ---\n{reference_pdf_text}"
            if reference_pdf_text else ""
        )
        header_text = f"--- REFERENCE DOCUMENT: {zname} ---{ref_section}\n\nChecklist items to evaluate against the document above:\n\n{items_block}"

        user_content: list = [{"type": "text", "text": header_text}]
        if mime == "pdf":
            cached_text, cached_pages = pdf_cache.get(zname, ("", []))
            if cached_text.strip():
                user_content.append({"type": "text", "text": f"Document text (for reference):{cached_text[:8000]}"})
            for page_png in cached_pages[:_PDF_MAX_PAGES]:
                b64 = base64.standard_b64encode(page_png).decode()
                user_content.append({"type": "image", "source": {"type": "base64", "media_type": "image/png", "data": b64}})
        else:
            guessed_mime = mimetypes.guess_type(zname)[0] or "image/png"
            safe = _claude_safe_image(fdata, guessed_mime)
            if safe is None:
                raise ValueError(f"Could not read image format: {zname} — try re-saving as JPG or PNG.")
            safe_data, safe_mime = safe
            b64 = base64.standard_b64encode(safe_data).decode()
            user_content.append({"type": "image", "source": {"type": "base64", "media_type": safe_mime, "data": b64}})
        return user_content

    def _run_batch(zname: str, idxs: list[int]) -> tuple[dict[int, dict], int, int]:
        try:
            user_content = _build_batch_user_content(zname, idxs)
            return _claude_check_batch(_claude_client, user_content, idxs)
        except Exception as e:
            fallback = {i: {"status": "N/A", "remark": f"Error analysing batched file: {e}"} for i in idxs}
            return fallback, 0, 0

    _active_jobs.add(uid)
    run_input_tokens = run_output_tokens = 0
    # One shared budget for this whole run: the batch pass and the per-item
    # pass below run one after the other, so a single semaphore keeps the
    # total in-flight Claude calls at CLAUDE_MAX_CONCURRENCY throughout.
    _claude_sem = asyncio.Semaphore(CLAUDE_MAX_CONCURRENCY)
    try:
        batch_results_by_index: dict[int, dict] = {}
        if batched_item_indices:
            batch_jobs = [
                (zname, idxs) for zname, idxs in batch_groups.items() if len(idxs) >= 2
            ]
            batch_outcomes = await asyncio.gather(
                *[_bounded_gather(_claude_sem, _run_batch, zname, idxs) for zname, idxs in batch_jobs]
            )
            for outcome, batch_in, batch_out in batch_outcomes:
                batch_results_by_index.update(outcome)
                run_input_tokens  += batch_in
                run_output_tokens += batch_out

        remaining = [(i, item) for i, item in non_section if i not in batched_item_indices]
        remaining_results = await asyncio.gather(
            *[_bounded_gather(_claude_sem, _analyse_item, i, item) for i, item in remaining]
        )
        for r in remaining_results:
            run_input_tokens  += r.get("_input_tokens", 0) or 0
            run_output_tokens += r.get("_output_tokens", 0) or 0
        results_by_index: dict[int, dict] = dict(zip((i for i, _ in remaining), remaining_results))
        results_by_index.update(batch_results_by_index)
        results = [results_by_index[i] for i, _ in non_section]
    finally:
        _active_jobs.discard(uid)

    # Fold in this quote's Verify Email token spend (if any) and consume it —
    # read from the DB's running total, not the submitted form fields, so a
    # stale/resubmitted page can't double-count or under-count against
    # whatever the DB actually accumulated across however many /email-verify
    # rounds happened. Reset to 0 so a future run for this quote starts fresh.
    if quote_id:
        try:
            qid_for_tokens = int(quote_id)
        except (ValueError, TypeError):
            qid_for_tokens = None
        if qid_for_tokens is not None:
            eml_in, eml_out = db.take_draft_email_verify_tokens(qid_for_tokens)
            run_input_tokens  += eml_in
            run_output_tokens += eml_out

    checklist_rows = []
    filled: dict[int, dict] = {}
    ns_iter  = iter(zip(non_section, results))
    next_ns  = next(ns_iter, None)
    for i, item in enumerate(items):
        if item.get("is_section"):
            checklist_rows.append({**item, "status": "", "remark": "", "ai_status": ""})
        else:
            (_, _item), result = next_ns
            filled[item["position"]] = result
            checklist_rows.append({**item, "status": result["status"], "remark": result["remark"],
                                    "ai_status": result["status"]})
            next_ns = next(ns_iter, None)

    # Parse email verify results passed from Step 2 (may be empty if step was skipped)
    email_results = []
    if email_results_json.strip():
        try:
            parsed = json.loads(email_results_json)
            if isinstance(parsed, list):
                email_results = [
                    r for r in parsed
                    if isinstance(r, dict) and r.get("name")
                ]
        except Exception:
            pass

    # Dedup exact duplicates: same sender address AND same attachment name
    # (same guard as the Step 2 upload handler) — keeps a single unified
    # answer per attachment even if results were re-submitted with repeats.
    if email_results:
        _seen_keys: set = set()
        _deduped = []
        for r in email_results:
            _key = (
                (r.get("source_address") or "").strip().lower(),
                (r.get("name") or "").strip().lower(),
            )
            if _key[0] and _key[1] and _key in _seen_keys:
                continue
            if _key[0] and _key[1]:
                _seen_keys.add(_key)
            _deduped.append(r)
        email_results = _deduped

    headers = {
        "customer_label": tpl.get("customer_label"),
        "address_label":  tpl.get("address_label"),
        "job_label":      tpl.get("job_label"),
    }
    xlsx_blob = build_xlsx(items, headers, tpl.get("note_text", ""), filled=filled,
                           email_results=email_results)

    if target_version_id is not None:
        # Admin "Upload ZIP" re-run against an existing Post-QC version —
        # overwrite it immediately (no draft/confirm review step; the whole
        # point of this action is "replace what's there"), then go straight
        # to the normal admin view of that same version.
        yes_count = sum(1 for r in checklist_rows if r.get("status") == "Yes")
        no_count  = sum(1 for r in checklist_rows if r.get("status") == "No")
        na_count  = sum(1 for r in checklist_rows if r.get("status") == "N/A")
        db.update_qc_version(
            version_id=target_version_id,
            xlsx_bytes=xlsx_blob,
            rows_json=json.dumps(checklist_rows),
            yes_count=yes_count,
            no_count=no_count,
            na_count=na_count,
            confirm=True,
            confirmed_by_user_id=uid,
            email_results_json="[]",
            input_tokens=run_input_tokens,
            output_tokens=run_output_tokens,
        )
        return RedirectResponse(url=f"/admin/qc-version/{target_version_id}", status_code=303)

    # Save Excel to a temp file on disk instead of embedding it in the DB row.
    # The signed token carries only the file path — a few bytes instead of megabytes.
    import tempfile
    xlsx_dir = os.path.join(os.path.dirname(__file__), "..", "tmp_xlsx")
    os.makedirs(xlsx_dir, exist_ok=True)
    xlsx_tmp = tempfile.NamedTemporaryFile(
        suffix=".xlsx", dir=xlsx_dir, delete=False
    )
    try:
        xlsx_tmp.write(xlsx_blob)
    finally:
        xlsx_tmp.close()

    dl_token   = _signer.dumps({"xlsx_path": xlsx_tmp.name, "name": tpl["name"]})
    tpl_safe   = {k: v for k, v in tpl.items() if not isinstance(v, (bytes, bytearray))}
    yes_count  = sum(1 for r in checklist_rows if r.get("status") == "Yes")
    no_count   = sum(1 for r in checklist_rows if r.get("status") == "No")
    na_count   = sum(1 for r in checklist_rows if r.get("status") == "N/A")
    result_token = store_pending_result({
        "tpl": tpl_safe,
        "rows": checklist_rows,
        "quote_id": quote_id,
        "dl_token": dl_token,
        "zip_filename": zip_file.filename or "",
        "yes_count": yes_count,
        "no_count": no_count,
        "na_count": na_count,
        "email_results": email_results,
        "input_tokens": run_input_tokens,
        "output_tokens": run_output_tokens,
    }, user_id=uid)
    return RedirectResponse(url=f"/checklist-result?token={result_token}", status_code=303)


# ---------------------------------------------------------------------------
# Checklist result
# ---------------------------------------------------------------------------

@router.get("/checklist-result", response_class=HTMLResponse)
def checklist_result(request: Request, token: str, user=Depends(require_qc_access)):
    payload = fetch_pending_result(token)
    if not payload:
        raise HTTPException(400, "Result link has expired. Please re-run the checklist.")

    preferred_install_date = ""
    quote_id = payload["quote_id"]
    record = None
    if quote_id:
        try:
            record = db.get_quote(int(quote_id))
        except (ValueError, TypeError):
            record = db.find_by_quote_number(quote_id)
        if record:
            preferred_install_date = record.get("preferred_install_date") or ""

    # Pull in the OTHER kind's latest saved version — same lookup
    # user_qc_version_revisit and admin_qc_version_view do for an
    # already-saved version. Without this, a fresh (not yet saved) run's
    # Confirm page showed "No Pre-QC run yet" even when a real Pre-QC
    # version already existed for this quote, just because this route
    # never looked it up.
    tpl_kind = (payload["tpl"] or {}).get("kind") or "pre"
    saved_email_results = payload.get("email_results", [])
    other_ctx = db.get_other_kind_context(quote_id, tpl_kind, saved_email_results)

    resp = templates.TemplateResponse(request, "user_result.html", {
        "current_user": user,
        "is_admin": False,
        "revisit_version": None,
        "tpl": payload["tpl"],
        "checklist_rows": payload["rows"],
        "dl_token": payload["dl_token"],
        "quote_id": quote_id,
        "zip_filename": payload.get("zip_filename", ""),
        "yes_count": payload.get("yes_count", 0),
        "no_count": payload.get("no_count", 0),
        "na_count": payload.get("na_count", 0),
        "email_results": saved_email_results,
        "preferred_install_date": preferred_install_date,
        "today": datetime.now().strftime("%Y-%m-%d"),
        "theme": _resolve_theme(user),
        "record": record,
        "input_tokens": payload.get("input_tokens", 0),
        "output_tokens": payload.get("output_tokens", 0),
        "claude_price_in": CLAUDE_PRICE_PER_M_INPUT,
        "claude_price_out": CLAUDE_PRICE_PER_M_OUTPUT,
        **other_ctx,
    })
    resp.headers.update(_NO_CACHE)
    return resp


@router.post("/checklist-analyse-file")
@limiter.limit("30/minute")
async def checklist_analyse_file(
    request: Request,
    files: list[UploadFile] = File(...),
    item_text: str = Form(""),
    prompt: str = Form(""),
    quote_id: str = Form(""),
    user=Depends(require_qc_access),
):
    """
    Analyse one or more uploaded files together against one checklist item's
    prompt — lets the user fix one row's evidence in edit mode without
    re-uploading the whole ZIP (which would re-run AI analysis, and cost, on
    every item). Independent of the manual Yes/No/N/A dropdown — neither
    depends on the other.
    """
    file_pairs: list[tuple[str, bytes]] = []
    total_bytes = 0
    for f in files:
        fdata = await f.read()
        total_bytes += len(fdata)
        if total_bytes > MAX_UPLOAD_BYTES:
            return {"status": "N/A", "remark": "Files are too large."}
        file_pairs.append((f.filename or "upload", fdata))

    reference_pdf_text = ""
    if quote_id:
        try:
            ref_record = db.get_quote(int(quote_id))
        except (ValueError, TypeError):
            ref_record = db.find_by_quote_number(quote_id)
        if ref_record:
            fields = [
                ("Quote Number",    ref_record.get("quote_number", "")),
                ("Customer Name",   ref_record.get("customer_name", "")),
                ("Install Date",    ref_record.get("install_date", "")),
                ("System Price",    ref_record.get("system_price", "")),
                ("STC Incentive",   ref_record.get("stc_incentive", "")),
                ("VIC Rebate",      ref_record.get("vic_rebate", "")),
                ("Battery Rebate",  ref_record.get("battery_rebate", "")),
                ("Total Price",     ref_record.get("total_price", "")),
                ("Deposit",         ref_record.get("deposit", "")),
                ("Balance",         ref_record.get("balance", "")),
                ("Payment Terms",   ref_record.get("payment_terms", "")),
                ("Billing Address", ref_record.get("billing_address", "")),
                ("Delivery Address",ref_record.get("delivery_address", "")),
                ("Roof Type",       ref_record.get("roof_type", "")),
                ("Email",           ref_record.get("email", "")),
                ("Phone",           ref_record.get("phone", "")),
                ("Notes",           ref_record.get("notes", "")),
            ]
            reference_pdf_text = "\n".join(f"{k}: {v}" for k, v in fields if v)
            line_items = ref_record.get("line_items") or []
            if line_items:
                reference_pdf_text += "\n\nSystem Components:\n" + "\n".join(
                    f"- {li.get('item','')}: {li.get('specification','')}"
                    for li in line_items if li.get("item")
                )

    client = _get_claude()
    result = _analyse_single_file(client, item_text, prompt, reference_pdf_text, file_pairs)
    return result


@router.post("/checklist-save-edits")
async def checklist_save_edits(request: Request, _auth=Depends(require_qc_access)):
    body       = await request.json()
    rows       = body.get("rows", [])
    tpl        = body.get("tpl", {})
    orig_token = body.get("dl_token", "")
    email_results = body.get("email_results", [])

    tpl_name = tpl.get("name", "checklist")
    prev_xlsx_path = None
    try:
        orig_payload = _signer.loads(orig_token, max_age=PENDING_TTL)
        tpl_name = orig_payload.get("name", tpl_name)
        # Preserve email_results from original token if not supplied by client
        if not email_results:
            email_results = orig_payload.get("email_results", [])
        # Remember the file this token pointed at so it can be removed once
        # its replacement is safely written (below) — every Save Changes
        # mints a new temp file, and without this the previous one was left
        # behind forever. A long editing session leaked one file per click.
        prev_xlsx_path = orig_payload.get("xlsx_path")
    except Exception:
        pass

    headers = {
        "customer_label": tpl.get("customer_label"),
        "address_label":  tpl.get("address_label"),
        "job_label":      tpl.get("job_label"),
    }
    note_text = tpl.get("note_text", "")

    filled = {}
    for row in rows:
        if not row.get("is_section") and row.get("position") is not None:
            filled[row["position"]] = {
                "status": row.get("status", "N/A"),
                "remark": row.get("remark", ""),
                "ai_status": row.get("ai_status", ""),
            }

    xlsx_blob = build_xlsx(rows, headers, note_text, filled=filled, email_results=email_results)

    import tempfile
    xlsx_dir = os.path.join(os.path.dirname(__file__), "..", "tmp_xlsx")
    os.makedirs(xlsx_dir, exist_ok=True)
    xlsx_tmp = tempfile.NamedTemporaryFile(
        suffix=".xlsx", dir=xlsx_dir, delete=False
    )
    try:
        xlsx_tmp.write(xlsx_blob)
    finally:
        xlsx_tmp.close()

    new_dl_token = _signer.dumps({"xlsx_path": xlsx_tmp.name, "name": tpl_name})

    # Only now that the replacement exists on disk is it safe to drop the
    # previous one — deleting first would leave a window where a crash mid-
    # write loses both. Guard the path so a forged/rewritten token can't
    # point this at an unrelated file outside tmp_xlsx/.
    if prev_xlsx_path and prev_xlsx_path != xlsx_tmp.name:
        try:
            if os.path.dirname(os.path.abspath(prev_xlsx_path)) == os.path.abspath(xlsx_dir):
                os.unlink(prev_xlsx_path)
        except OSError:
            pass

    return {"dl_token": new_dl_token}


@router.get("/checklist-download", response_class=Response)
def checklist_download(token: str, _auth=Depends(require_qc_access)):
    try:
        payload = _signer.loads(token, max_age=PENDING_TTL)
    except BadSignature:
        raise HTTPException(400, "Invalid or expired download token.")

    # Support both old base64 tokens and new file-path tokens
    if "xlsx_path" in payload:
        xlsx_path = payload["xlsx_path"]
        if not os.path.isfile(xlsx_path):
            raise HTTPException(410, "Download file has expired. Please re-run the checklist.")
        with open(xlsx_path, "rb") as f:
            xlsx_blob = f.read()
    else:
        xlsx_blob = base64.b64decode(payload["xlsx"])

    safe_name = (payload.get("name") or "checklist").replace(" ", "_")
    return Response(
        content=xlsx_blob,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={
            "Content-Disposition": f'attachment; filename="{safe_name}_QC.xlsx"',
            "Cache-Control": "no-store, no-cache, must-revalidate",
            "Pragma": "no-cache",
        },
    )


@router.post("/checklist-confirm")
def checklist_confirm(
    request: Request,
    quote_id: str = Form(""),
    dl_token: str = Form(""),
    zip_filename: str = Form(""),
    tpl_name: str = Form(""),
    tpl_kind: str = Form(""),
    yes_count: int = Form(0),
    no_count: int = Form(0),
    na_count: int = Form(0),
    rows_json: str = Form(""),
    email_results_json: str = Form(""),
    preferred_install_date: str = Form(""),
    version_id: str = Form(""),
    input_tokens: int = Form(0),
    output_tokens: int = Form(0),
    user=Depends(require_qc_access),
):
    tpl_name     = tpl_name[:200].replace("\n", "").replace("\r", "")
    zip_filename = zip_filename[:260].replace("\n", "").replace("\r", "")
    kind = tpl_kind if tpl_kind in ("pre", "post") else "pre"
    # Post-QC never legitimately collects its own email verification data —
    # refuse to persist any into a post-kind version, closing the same hole
    # user_qc_version_save guards against.
    if kind == "post":
        email_results_json = "[]"

    record = None
    if quote_id:
        try:
            record = db.get_quote(int(quote_id))
        except (ValueError, TypeError):
            record = db.find_by_quote_number(quote_id)

    if record and preferred_install_date.strip() != (record.get("preferred_install_date") or ""):
        db.update_preferred_install_date(record["id"], preferred_install_date.strip())
        record["preferred_install_date"] = preferred_install_date.strip()

    if dl_token and record:
        try:
            payload = _signer.loads(dl_token, max_age=PENDING_TTL)
            if "xlsx_path" in payload:
                xlsx_path = payload["xlsx_path"]
                try:
                    with open(xlsx_path, "rb") as f:
                        xlsx_bytes = f.read()
                    # Clean up the temp file now that it's been saved to DB
                    try:
                        os.unlink(xlsx_path)
                    except OSError:
                        pass
                except FileNotFoundError:
                    # See the matching comment in checklist_save_draft — the
                    # temp xlsx this token pointed at was already cleaned up
                    # (2-hour TTL), most likely a revisit page left open past
                    # that window. Rebuild from rows_json instead of failing.
                    logger.warning(
                        "checklist_confirm: tmp xlsx missing at %s, rebuilding from rows_json (quote_id=%s)",
                        xlsx_path, quote_id,
                    )
                    _rebuild_rows = json.loads(rows_json) if rows_json else []
                    _rebuild_filled = {}
                    for _r in _rebuild_rows:
                        if not _r.get("is_section") and _r.get("position") is not None:
                            _rebuild_filled[_r["position"]] = {
                                "status": _r.get("status", "N/A"),
                                "remark": _r.get("remark", ""),
                                "ai_status": _r.get("ai_status", ""),
                            }
                    xlsx_bytes = build_xlsx(
                        _rebuild_rows, filled=_rebuild_filled,
                        email_results=json.loads(email_results_json) if email_results_json else [],
                    )
            else:
                xlsx_bytes = base64.b64decode(payload["xlsx"])
            db.save_qc_excel(record["id"], xlsx_bytes)
            # A fresh run (no version_id — the field only carries a value on
            # an explicit revisit/re-run) can still be for a quote+kind that
            # already has a version, e.g. running Post-QC again from scratch
            # for a customer already confirmed or draft-saved earlier. Reuse
            # that version instead of creating a duplicate — same as an
            # explicit revisit — so History can't accumulate multiple rows
            # for what a technician experiences as "the same QC run".
            target_version_id = int(version_id) if version_id.strip() else None
            if target_version_id is None:
                existing = db.get_latest_qc_version(record["id"], kind)
                if existing:
                    target_version_id = existing["id"]
            if target_version_id is not None:
                db.update_qc_version(
                    version_id=target_version_id,
                    xlsx_bytes=xlsx_bytes,
                    rows_json=rows_json,
                    yes_count=yes_count,
                    no_count=no_count,
                    na_count=na_count,
                    confirm=True,
                    confirmed_by_user_id=user["id"],
                    email_results_json=email_results_json,
                    input_tokens=input_tokens,
                    output_tokens=output_tokens,
                )
            else:
                db.add_qc_version(
                    quote_id=record["id"],
                    xlsx_bytes=xlsx_bytes,
                    template_name=tpl_name,
                    zip_filename=zip_filename,
                    yes_count=yes_count,
                    no_count=no_count,
                    na_count=na_count,
                    rows_json=rows_json,
                    email_results_json=email_results_json,
                    confirmed_by_user_id=user["id"],
                    saved_by_user_id=user["id"],
                    status="confirmed",
                    kind=kind,
                    input_tokens=input_tokens,
                    output_tokens=output_tokens,
                )
        except Exception:
            logger.exception("checklist_confirm failed to save quote_id=%s", quote_id)
            resp = templates.TemplateResponse(request, "user_home.html", _index_context(
                user=user,
                error="Something went wrong saving this QC result. Nothing was lost — please try Confirm again.",
            ))
            resp.headers.update(_NO_CACHE)
            return resp

    preferred = (record or {}).get("preferred_install_date", "").strip() if record else ""
    install_date = preferred or ((record or {}).get("install_date", "").strip() if record else "")
    cal_param = ""
    if install_date:
        for fmt in ("%Y-%m-%d", "%d/%m/%Y", "%m/%d/%Y", "%d-%m-%Y"):
            try:
                dt = datetime.strptime(install_date, fmt)
                cal_param = f"?cal={dt.year}-{dt.month:02d}"
                break
            except ValueError:
                continue

    resp = templates.TemplateResponse(request, "user_thankyou.html", {
        "current_user": user,
        "theme": _resolve_theme(user),
        "cal_param": cal_param,
    })
    resp.headers.update(_NO_CACHE)
    return resp


@router.post("/checklist-save-draft")
def checklist_save_draft(
    request: Request,
    quote_id: str = Form(""),
    dl_token: str = Form(""),
    zip_filename: str = Form(""),
    tpl_name: str = Form(""),
    tpl_kind: str = Form(""),
    yes_count: int = Form(0),
    no_count: int = Form(0),
    na_count: int = Form(0),
    rows_json: str = Form(""),
    email_results_json: str = Form(""),
    preferred_install_date: str = Form(""),
    version_id: str = Form(""),
    input_tokens: int = Form(0),
    output_tokens: int = Form(0),
    user=Depends(require_qc_access),
):
    tpl_name     = tpl_name[:200].replace("\n", "").replace("\r", "")
    zip_filename = zip_filename[:260].replace("\n", "").replace("\r", "")
    kind = tpl_kind if tpl_kind in ("pre", "post") else "pre"
    # Post-QC never legitimately collects its own email verification data —
    # refuse to persist any into a post-kind version, closing the same hole
    # user_qc_version_save guards against.
    if kind == "post":
        email_results_json = "[]"

    record = None
    if quote_id:
        try:
            record = db.get_quote(int(quote_id))
        except (ValueError, TypeError):
            record = db.find_by_quote_number(quote_id)

    if record and preferred_install_date.strip() != (record.get("preferred_install_date") or ""):
        db.update_preferred_install_date(record["id"], preferred_install_date.strip())

    if dl_token and record:
        try:
            payload = _signer.loads(dl_token, max_age=PENDING_TTL)
            if "xlsx_path" in payload:
                xlsx_path = payload["xlsx_path"]
                try:
                    with open(xlsx_path, "rb") as f:
                        xlsx_bytes = f.read()
                    try:
                        os.unlink(xlsx_path)
                    except OSError:
                        pass
                except FileNotFoundError:
                    # The temp xlsx this token pointed at was already cleaned
                    # up (2-hour TTL, same as the token's own max_age) — e.g.
                    # a revisit page left open in a tab past that window. The
                    # actual checklist data is still right here in rows_json,
                    # so rebuild the Excel from that instead of failing the
                    # whole save over a missing temp file.
                    logger.warning(
                        "checklist_save_draft: tmp xlsx missing at %s, rebuilding from rows_json (quote_id=%s)",
                        xlsx_path, quote_id,
                    )
                    _rebuild_rows = json.loads(rows_json) if rows_json else []
                    _rebuild_filled = {}
                    for _r in _rebuild_rows:
                        if not _r.get("is_section") and _r.get("position") is not None:
                            _rebuild_filled[_r["position"]] = {
                                "status": _r.get("status", "N/A"),
                                "remark": _r.get("remark", ""),
                                "ai_status": _r.get("ai_status", ""),
                            }
                    xlsx_bytes = build_xlsx(
                        _rebuild_rows, filled=_rebuild_filled,
                        email_results=json.loads(email_results_json) if email_results_json else [],
                    )
            else:
                xlsx_bytes = base64.b64decode(payload["xlsx"])
            # Same reuse-if-one-already-exists rule as checklist_confirm —
            # see its comment for why. Save Draft is the more common repeat
            # offender in practice (a technician saving progress multiple
            # times on the same customer before finally confirming).
            #
            # Only reuse an existing DRAFT, never a CONFIRMED one — Save
            # Draft always writes status='draft', and reusing a confirmed
            # version here would silently demote a signed-off result back
            # to draft. A fresh Save Draft on top of an already-confirmed
            # customer instead creates a new draft alongside it, leaving
            # the confirmed version untouched.
            target_version_id = int(version_id) if version_id.strip() else None
            if target_version_id is None:
                existing = db.get_latest_qc_version(record["id"], kind)
                if existing and existing.get("status") == "draft":
                    target_version_id = existing["id"]
            if target_version_id is not None:
                db.update_qc_version(
                    version_id=target_version_id,
                    xlsx_bytes=xlsx_bytes,
                    rows_json=rows_json,
                    yes_count=yes_count,
                    no_count=no_count,
                    na_count=na_count,
                    confirm=False,
                    email_results_json=email_results_json,
                    input_tokens=input_tokens,
                    output_tokens=output_tokens,
                )
            else:
                db.add_qc_version(
                    quote_id=record["id"],
                    xlsx_bytes=xlsx_bytes,
                    template_name=tpl_name,
                    zip_filename=zip_filename,
                    yes_count=yes_count,
                    no_count=no_count,
                    na_count=na_count,
                    rows_json=rows_json,
                    email_results_json=email_results_json,
                    saved_by_user_id=user["id"],
                    status="draft",
                    kind=kind,
                    input_tokens=input_tokens,
                    output_tokens=output_tokens,
                )
        except Exception:
            logger.exception("checklist_save_draft failed to save quote_id=%s", quote_id)
            resp = templates.TemplateResponse(request, "user_home.html", _index_context(
                user=user,
                error="Something went wrong saving this draft. Nothing was lost — please try Save Draft again.",
            ))
            resp.headers.update(_NO_CACHE)
            return resp

    return RedirectResponse(url="/user/history", status_code=303)


# ---------------------------------------------------------------------------
# User history
# ---------------------------------------------------------------------------

@router.get("/user/history", response_class=HTMLResponse)
def user_history(request: Request, user=Depends(require_qc_access)):
    history = db.get_qc_history_for_user(user["id"])
    resp = templates.TemplateResponse(request, "user_history.html", {
        "current_user": user,
        "theme": _resolve_theme(user),
        "history": history,
    })
    resp.headers.update(_NO_CACHE)
    return resp


@router.get("/user/qc-version/{version_id}", response_class=HTMLResponse)
def user_qc_version_revisit(request: Request, version_id: int, user=Depends(require_qc_access)):
    with get_db() as conn:
        row = conn.execute(
            """SELECT qv.*, q.customer_name, q.quote_number, q.install_date,
                      q.preferred_install_date
               FROM qc_versions qv JOIN quotes q ON q.id = qv.quote_id
               WHERE qv.id = ?""",
            (version_id,),
        ).fetchone()
    if not row:
        raise HTTPException(404, "QC version not found.")
    v = dict(row)
    is_own_version = (
        v.get("saved_by_user_id") == user["id"]
        or v.get("confirmed_by_user_id") == user["id"]
    )
    if user.get("role") != "admin":
        # A user assigned to this customer for Post-QC has already been
        # vetted to see this customer's QC work (see post_qc_customer_detail)
        # — let them click through to ANY version for that customer (Pre-QC
        # included), not just runs they personally confirmed, so the QC
        # History list on their assigned-customer page is actually clickable.
        assignee = db.get_post_qc_assignee(v["quote_id"])
        is_assigned_to_customer = assignee is not None and assignee.get("user_id") == user["id"]
        if not is_own_version and not is_assigned_to_customer:
            raise HTTPException(403, "Access denied.")
    # Viewing someone else's version (reached only via the customer-assignment
    # path above, or an admin browsing) is read-only — the save endpoint
    # already rejects this, but the page shouldn't show Modify/Save/Confirm
    # controls that would just fail, especially for a role like Post-QC that
    # should never edit another person's Pre-QC work.
    can_edit = user.get("role") == "admin" or is_own_version
    try:
        rows = json.loads(v["rows_json"]) if v.get("rows_json") else []
        if not isinstance(rows, list):
            rows = []
    except Exception:
        rows = []

    tpl = {
        "id": 0, "name": v.get("template_name", ""),
        "customer_label": "", "address_label": "", "job_label": "", "note_text": "",
        "kind": v.get("kind") or "pre",
    }
    # For revisiting saved versions, write the Excel back to a temp file so the
    # dl_token stays small (same pattern as the live checklist flow).
    # get_qc_version_excel is disk-first with a legacy-blob fallback.
    dl_token = ""
    _xlsx_bytes = db.get_qc_version_excel(version_id)
    if _xlsx_bytes:
        import tempfile
        xlsx_dir = os.path.join(os.path.dirname(__file__), "..", "tmp_xlsx")
        os.makedirs(xlsx_dir, exist_ok=True)
        xlsx_tmp = tempfile.NamedTemporaryFile(
            suffix=".xlsx", dir=xlsx_dir, delete=False
        )
        try:
            xlsx_tmp.write(_xlsx_bytes)
        finally:
            xlsx_tmp.close()
        dl_token = _signer.dumps({"xlsx_path": xlsx_tmp.name, "name": tpl["name"]})

    yes_count = sum(1 for r in rows if not r.get("is_section") and r.get("status") == "Yes")
    no_count  = sum(1 for r in rows if not r.get("is_section") and r.get("status") == "No")
    na_count  = sum(1 for r in rows if not r.get("is_section") and r.get("status") == "N/A")

    # Load saved email results from DB
    saved_email_results = []
    raw_email_json = v.get("email_results_json") or ""
    if raw_email_json.strip():
        try:
            parsed_email = json.loads(raw_email_json)
            if isinstance(parsed_email, list):
                saved_email_results = parsed_email
        except Exception:
            pass

    # Full PDF-extracted quote record for the Signed Agreement tile/modal —
    # same data admin_qc_version_view fetches for its own equivalent block.
    record = db.get_quote(v["quote_id"])

    # Pull in the OTHER kind's latest version too — same as admin's view, so
    # a Post-QC user can see that customer's Pre-QC results (read-only, via
    # its own view-only page) even though someone else ran it, instead of
    # the tile just sitting empty.
    other_ctx = db.get_other_kind_context(v["quote_id"], tpl["kind"], saved_email_results)

    resp = templates.TemplateResponse(request, "user_result.html", {
        "current_user": user,
        "theme": _resolve_theme(user),
        "tpl": tpl,
        "checklist_rows": rows,
        "dl_token": dl_token,
        "quote_id": v["quote_id"],
        "zip_filename": v.get("zip_filename", ""),
        "yes_count": yes_count,
        "no_count": no_count,
        "na_count": na_count,
        "email_results": saved_email_results,
        "revisit_version": v,
        "preferred_install_date": v.get("preferred_install_date") or "",
        "today": datetime.now().strftime("%Y-%m-%d"),
        "record": record,
        "can_edit": can_edit,
        **other_ctx,
    })
    resp.headers.update(_NO_CACHE)
    return resp


@router.post("/user/qc-version/{version_id}/save")
async def user_qc_version_save(request: Request, version_id: int, user=Depends(require_qc_access)):
    """Let the user who saved/confirmed a QC version edit it in place —
    same action the admin panel already offers on any version. Overwrites
    the existing version row rather than creating a new one."""
    with get_db() as conn:
        row = conn.execute(
            "SELECT * FROM qc_versions WHERE id = ?", (version_id,)
        ).fetchone()
    if not row:
        raise HTTPException(404, "QC version not found.")
    v = dict(row)
    if (user.get("role") != "admin"
            and v.get("saved_by_user_id") != user["id"]
            and v.get("confirmed_by_user_id") != user["id"]):
        raise HTTPException(403, "Access denied.")

    body = await request.json()
    rows = body.get("rows", [])
    preferred_install_date = (body.get("preferred_install_date") or "").strip()
    # Per-row re-analysis (analyseRowFile) while editing costs real Claude
    # tokens on top of whatever this version already recorded — the browser
    # sends the delta from just this edit session, added onto the version's
    # existing stored total (never trust a client-sent absolute total, only
    # an incremental delta on top of what the server already has on file).
    input_tokens_delta  = int(body.get("input_tokens_delta", 0) or 0)
    output_tokens_delta = int(body.get("output_tokens_delta", 0) or 0)
    new_input_tokens  = (v.get("input_tokens")  or 0) + input_tokens_delta
    new_output_tokens = (v.get("output_tokens") or 0) + output_tokens_delta
    # action is optional — omitted means "keep whatever status this version
    # already had" (the original revisit-save behavior); "draft" or "confirm"
    # explicitly sets the status, same choice the admin panel already offers.
    action = body.get("action")
    if action not in ("draft", "confirm"):
        action = "confirm" if v.get("status") == "confirmed" else "draft"
    # email_results is optional in the payload — omitted (key absent) means
    # "the email table wasn't part of this edit, leave it as stored"; an
    # empty list is a legitimate value (all rows cleared) and must overwrite.
    has_email_edit = "email_results" in body
    email_results = body.get("email_results", [])
    # Post-QC never legitimately collects its own email verification data —
    # the client only ever shows email rows here when they belong to THIS
    # version. But the Email Verification tile can display Pre-QC's data
    # borrowed onto a Post-QC page (read-only in the UI); refuse to let a
    # Post-QC save persist non-empty email results into its own row, so a
    # stale client or future UI bug can't re-create that corruption server-side.
    if (v.get("kind") or "pre") == "post" and email_results:
        email_results = []

    try:
        filled = {}
        for r in rows:
            if not r.get("is_section") and r.get("position") is not None:
                filled[r["position"]] = {
                    "status": r.get("status", "N/A"),
                    "remark": r.get("remark", ""),
                    "ai_status": r.get("ai_status", ""),
                }

        xlsx_blob = build_xlsx(
            rows, filled=filled,
            email_results=email_results if has_email_edit else json.loads(v.get("email_results_json") or "[]"),
        )
        yes_count = sum(1 for r in rows if not r.get("is_section") and r.get("status") == "Yes")
        no_count  = sum(1 for r in rows if not r.get("is_section") and r.get("status") == "No")
        na_count  = sum(1 for r in rows if not r.get("is_section") and r.get("status") == "N/A")

        confirm = action == "confirm"
        db.update_qc_version(
            version_id=version_id,
            xlsx_bytes=xlsx_blob,
            rows_json=json.dumps(rows),
            yes_count=yes_count,
            no_count=no_count,
            na_count=na_count,
            confirm=confirm,
            confirmed_by_user_id=user["id"] if confirm else None,
            email_results_json=json.dumps(email_results) if has_email_edit else None,
            input_tokens=new_input_tokens if (input_tokens_delta or output_tokens_delta) else None,
            output_tokens=new_output_tokens if (input_tokens_delta or output_tokens_delta) else None,
        )

        quote_id = v["quote_id"]
        record = db.get_quote(quote_id)
        if record and preferred_install_date != (record.get("preferred_install_date") or ""):
            db.update_preferred_install_date(quote_id, preferred_install_date)
    except Exception as e:
        raise HTTPException(500, f"Failed to save: {e}")

    return {"ok": True, "yes_count": yes_count, "no_count": no_count, "na_count": na_count}


@router.post("/user/qc-version/{version_id}/delete")
def user_qc_version_delete(request: Request, version_id: int, user=Depends(require_qc_access)):
    """Let the user who saved/confirmed a QC version delete it themselves —
    same soft-delete admin_delete_qc_version already offers, same ownership
    rule user_qc_version_save uses (own it, or be admin). A user deleting
    their own mistaken/duplicate run shows up in the admin trash immediately
    since it's the exact same qc_versions row and is_deleted flag."""
    with get_db() as conn:
        row = conn.execute(
            """SELECT qv.quote_id, qv.version, qv.template_name, COALESCE(qv.kind, 'pre') as kind,
                      qv.saved_by_user_id, qv.confirmed_by_user_id,
                      q.quote_number, q.customer_name
               FROM qc_versions qv JOIN quotes q ON q.id = qv.quote_id
               WHERE qv.id = ?""",
            (version_id,),
        ).fetchone()
    if not row:
        raise HTTPException(404, "QC version not found.")
    v = dict(row)
    if (user.get("role") != "admin"
            and v.get("saved_by_user_id") != user["id"]
            and v.get("confirmed_by_user_id") != user["id"]):
        raise HTTPException(403, "Access denied.")

    db.soft_delete_qc_version(version_id, deleted_by_user_id=user["id"])
    audit.log_event(request, "qc_version_deleted", username=user["username"], user_id=user["id"],
                    detail=f"{v['kind']}-QC v{v['version']} ({v.get('template_name') or ''}) for quote "
                           f"{v.get('quote_number') or v['quote_id']} ({v.get('customer_name') or ''})")
    return RedirectResponse(url="/user/history", status_code=303)


@router.post("/user/qc-version/{version_id}/save-email-only")
async def user_qc_version_save_email_only(request: Request, version_id: int, user=Depends(require_qc_access)):
    """Edit the email_results on a QC version WITHOUT touching its checklist
    rows — used from the sibling-kind page (e.g. editing Pre-QC's email
    results while viewing the Post-QC page that borrows them read-only,
    since Post-QC never collects its own). Unlike /save (which rebuilds the
    Excel from whatever "rows" the browser currently has displayed — the
    OTHER kind's checklist on that page, which would silently corrupt this
    version's real checklist data), this always reloads the target
    version's OWN stored rows_json from the DB and rebuilds its Excel with
    those unchanged, only swapping in the new email_results."""
    with get_db() as conn:
        row = conn.execute(
            "SELECT * FROM qc_versions WHERE id = ?", (version_id,)
        ).fetchone()
    if not row:
        raise HTTPException(404, "QC version not found.")
    v = dict(row)
    if (user.get("role") != "admin"
            and v.get("saved_by_user_id") != user["id"]
            and v.get("confirmed_by_user_id") != user["id"]):
        raise HTTPException(403, "Access denied.")
    # Post-QC versions never legitimately hold their own email results (same
    # rule /save enforces) — this endpoint only ever makes sense for editing
    # a Pre-QC version's real data, so refuse outright rather than silently
    # writing into a Post-QC row that should always stay empty.
    if (v.get("kind") or "pre") == "post":
        raise HTTPException(400, "Post-QC versions do not hold their own email verification data.")

    body = await request.json()
    email_results = body.get("email_results", [])

    try:
        own_rows = json.loads(v.get("rows_json") or "[]")
        if not isinstance(own_rows, list):
            own_rows = []
        filled = {}
        for r in own_rows:
            if not r.get("is_section") and r.get("position") is not None:
                filled[r["position"]] = {
                    "status": r.get("status", "N/A"),
                    "remark": r.get("remark", ""),
                    "ai_status": r.get("ai_status", ""),
                }
        xlsx_blob = build_xlsx(own_rows, filled=filled, email_results=email_results)
        yes_count = sum(1 for r in own_rows if not r.get("is_section") and r.get("status") == "Yes")
        no_count  = sum(1 for r in own_rows if not r.get("is_section") and r.get("status") == "No")
        na_count  = sum(1 for r in own_rows if not r.get("is_section") and r.get("status") == "N/A")
        db.update_qc_version(
            version_id=version_id,
            xlsx_bytes=xlsx_blob,
            rows_json=json.dumps(own_rows),
            yes_count=yes_count,
            no_count=no_count,
            na_count=na_count,
            confirm=(v.get("status") == "confirmed"),
            confirmed_by_user_id=v.get("confirmed_by_user_id"),
            email_results_json=json.dumps(email_results),
        )
    except Exception as e:
        raise HTTPException(500, f"Failed to save: {e}")

    return {"ok": True}


@router.post("/user/qc-version/{version_id}/add-email")
@limiter.limit("15/minute")
async def user_qc_version_add_email(
    request: Request,
    version_id: int,
    eml_file: UploadFile = File(...),
    user=Depends(require_qc_access),
):
    """Analyse one more .eml file against the signed agreement and append its
    attachment results to this version's saved email results immediately —
    independent of the Modify/edit-mode flow, no separate Save step."""
    with get_db() as conn:
        row = conn.execute(
            "SELECT * FROM qc_versions WHERE id = ?", (version_id,)
        ).fetchone()
    if not row:
        raise HTTPException(404, "QC version not found.")
    v = dict(row)
    if (user.get("role") != "admin"
            and v.get("saved_by_user_id") != user["id"]
            and v.get("confirmed_by_user_id") != user["id"]):
        raise HTTPException(403, "Access denied.")

    if not eml_file.filename.lower().endswith(".eml"):
        raise HTTPException(400, f"'{eml_file.filename}' is not a valid .eml file (Outlook email format).")

    raw = await eml_file.read()
    if len(raw) > 50 * 1024 * 1024:
        raise HTTPException(400, f"'{eml_file.filename}' is too large. Maximum allowed size is 50 MB.")

    import email as _email_mod
    import email.policy as _policy

    msg = _email_mod.message_from_bytes(raw, policy=_policy.compat32)
    subject = _decode_email_field(msg.get("Subject", "")) or "(no subject)"
    from_hdr = _decode_email_field(msg.get("From", "")) or "(unknown sender)"
    address = _extract_email_address(from_hdr)

    attachments = _parse_eml(raw)
    if not attachments:
        raise HTTPException(400, f"No attachments found in '{eml_file.filename}'. Please check the .eml file.")

    reference_text = _build_reference_text(v["quote_id"])
    client = _get_claude()

    _eml_sem = asyncio.Semaphore(CLAUDE_MAX_CONCURRENCY)
    analysed = await asyncio.gather(
        *[_bounded_gather(_eml_sem, _match_attachment_with_claude, client, att, reference_text)
          for att in attachments]
    )

    new_results = [
        {
            **r,
            "ai_status": r.get("status"),
            "source_email": from_hdr,
            "source_subject": subject,
            "source_address": address,
        }
        for r in analysed
    ]

    try:
        existing_results = json.loads(v.get("email_results_json") or "[]") or []
    except Exception:
        existing_results = []

    new_keys = {
        ((r.get("source_address") or ""), (r.get("name") or "").strip().lower())
        for r in new_results
    }
    kept_old = [
        r for r in existing_results
        if ((r.get("source_address") or ""), (r.get("name") or "").strip().lower())
        not in new_keys
    ]
    merged_results = kept_old + new_results

    rows = json.loads(v["rows_json"]) if v.get("rows_json") else []
    filled = {}
    for r in rows:
        if not r.get("is_section") and r.get("position") is not None:
            filled[r["position"]] = {
                "status": r.get("status", "N/A"),
                "remark": r.get("remark", ""),
                "ai_status": r.get("ai_status", ""),
            }
    xlsx_blob = build_xlsx(rows, filled=filled, email_results=merged_results)

    was_confirmed = v.get("status") == "confirmed"
    db.update_qc_version(
        version_id=version_id,
        xlsx_bytes=xlsx_blob,
        rows_json=v["rows_json"] or json.dumps(rows),
        yes_count=v["yes_count"],
        no_count=v["no_count"],
        na_count=v["na_count"],
        confirm=was_confirmed,
        confirmed_by_user_id=user["id"] if was_confirmed else None,
        email_results_json=json.dumps(merged_results),
    )
    return {"ok": True, "email_results": merged_results}


# ---------------------------------------------------------------------------
# Post-QC — assigned-customer entry point
# ---------------------------------------------------------------------------

@router.get("/post-qc", response_class=HTMLResponse)
def post_qc_customers(request: Request, user=Depends(require_qc_access)):
    """List Post-Install QC work to do — the Post-QC equivalent of Step 1
    (Upload PDF): no new PDF is uploaded here, since the reference/quote
    data already exists from that customer's Pre-QC run.

    Admin/super-admin see EVERY customer still pending Post-QC (they have
    blanket access, not per-customer assignments, so the assignment list
    would always be empty for them). Regular Post-QC technicians keep
    seeing only the customers explicitly assigned to them."""
    if user.get("role") != "admin" and not user.get("can_post_qc"):
        raise HTTPException(403, "You do not have Post-QC access.")
    is_admin = user.get("role") == "admin"
    assigned = db.get_quotes_pending_post_qc() if is_admin else db.get_assigned_quotes_for_user(user["id"])
    resp = templates.TemplateResponse(request, "user_post_qc.html", {
        "current_user": user,
        "theme": _resolve_theme(user),
        "assigned": assigned,
        "is_admin_view": is_admin,
    })
    resp.headers.update(_NO_CACHE)
    return resp


@router.get("/post-qc/customer/{quote_id}", response_class=HTMLResponse)
def post_qc_customer_detail(request: Request, quote_id: int, user=Depends(require_qc_access)):
    """Read-only view of an assigned customer's existing reference/quote
    details, with a link into the ZIP-upload step (Post-QC's own Step 2),
    filtered to Post-QC templates only."""
    if user.get("role") != "admin" and not user.get("can_post_qc"):
        raise HTTPException(403, "You do not have Post-QC access.")
    record = db.get_quote(quote_id)
    if not record:
        raise HTTPException(404, "Customer record not found.")
    if user.get("role") != "admin":
        assignee = db.get_post_qc_assignee(quote_id)
        if not assignee or assignee.get("user_id") != user["id"]:
            raise HTTPException(403, "This customer is not assigned to you for Post-QC.")
    qc_versions = db.get_qc_versions(quote_id)
    resp = templates.TemplateResponse(request, "user_post_qc_customer.html", {
        "current_user": user,
        "theme": _resolve_theme(user),
        "record": record,
        "qc_versions": qc_versions,
    })
    resp.headers.update(_NO_CACHE)
    return resp
